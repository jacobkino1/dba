import os
import re
import uuid
from pathlib import Path
from typing import List, Optional
from datetime import datetime
from auth_models import (
    LoginRequest,
    WorkstationLoginRequest,
    LoginResponse,
    ChangePasswordRequest,
    ResetPasswordRequest,
    AuthUserPayload,
    ClinicMembershipPayload,
)

from dotenv import load_dotenv
load_dotenv()
from auth_utils import create_access_token, verify_password, hash_password

from database.db import engine
from database.models import Base
from sqlalchemy.orm import Session
from sqlalchemy import desc
from database.db import SessionLocal
from database.models import (
    Organisation,
    Clinic,
    Document,
    DocumentAuditLog,
    User,
    OrganisationMembership,
    ClinicMembership,
    NetworkAccess,
    AllowedIP,
)
from security import CurrentUser, get_current_user, get_client_ip, require_organisation_access, require_selected_clinic
from permissions import (
    VALID_PERMISSION_LEVELS,
    normalize_permission_level,
    has_level,
    can_view_document,
    can_download_document,
    can_upload_document,
    can_replace_document,
    can_archive_document,
    can_restore_document,
    can_delete_document,
    can_view_activity,
    can_create_users,
    can_grant_level,
    can_grant_document_level,
    can_assign_org_scope,
    can_manage_shared_documents,
    get_accessible_document_levels,
)

from fastapi import FastAPI, UploadFile, File, Depends, BackgroundTasks, Form, HTTPException, Request
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from pydantic import Field

from chat_routes import router as chat_router
from ask_logging import create_ask_log
from insights_routes import router as insights_router

from qdrant_client import QdrantClient
from qdrant_client.http import models as qm

from pypdf import PdfReader
from docx import Document as DocxDocument
from openai import AzureOpenAI
import random

AZURE_OPENAI_ENDPOINT = os.getenv("AZURE_OPENAI_ENDPOINT")
AZURE_OPENAI_API_KEY = os.getenv("AZURE_OPENAI_API_KEY")
AZURE_OPENAI_API_VERSION = os.getenv("AZURE_OPENAI_API_VERSION")

CHAT_DEPLOYMENT = os.getenv("AZURE_OPENAI_CHAT_DEPLOYMENT")
EMBEDDING_DEPLOYMENT = os.getenv("AZURE_OPENAI_EMBEDDING_DEPLOYMENT")

azure_client = AzureOpenAI(
    api_key=AZURE_OPENAI_API_KEY,
    api_version=AZURE_OPENAI_API_VERSION,
    azure_endpoint=AZURE_OPENAI_ENDPOINT,
)

Base.metadata.create_all(bind=engine)

def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

def create_document_audit_log(
    db: Session,
    *,
    action: str,
    organisation_id: str,
    clinic_id: str | None,
    performed_by: str | None,
    filename: str | None = None,
    document_id: str | None = None,
    old_document_id: str | None = None,
    new_document_id: str | None = None,
    notes: str | None = None,
):
    audit_log = DocumentAuditLog(
        auditId=str(uuid.uuid4()),
        documentId=document_id,
        oldDocumentId=old_document_id,
        newDocumentId=new_document_id,
        organisationId=organisation_id,
        clinicId=clinic_id,
        action=action,
        performedBy=performed_by,
        filename=filename,
        notes=notes,
    )
    db.add(audit_log)

def get_user_by_id(db: Session, user_id: str) -> User | None:
    return db.query(User).filter(User.userId == user_id).first()

def get_organisation_membership(
    db: Session,
    *,
    user_id: str,
    organisation_id: str,
) -> OrganisationMembership | None:
    return (
        db.query(OrganisationMembership)
        .filter(
            OrganisationMembership.userId == user_id,
            OrganisationMembership.organisationId == organisation_id,
        )
        .first()
    )

def get_clinic_membership(
    db: Session,
    *,
    user_id: str,
    organisation_id: str,
    clinic_id: str,
) -> ClinicMembership | None:
    return (
        db.query(ClinicMembership)
        .filter(
            ClinicMembership.userId == user_id,
            ClinicMembership.organisationId == organisation_id,
            ClinicMembership.clinicId == clinic_id,
        )
        .first()
    )

def get_clinic_memberships_for_user(
    db: Session,
    *,
    user_id: str,
    organisation_id: str,
) -> list[ClinicMembership]:
    return (
        db.query(ClinicMembership)
        .filter(
            ClinicMembership.userId == user_id,
            ClinicMembership.organisationId == organisation_id,
        )
        .all()
    )

def get_user_by_email(db: Session, email: str) -> User | None:
    normalized_email = email.strip().lower()
    return db.query(User).filter(User.email == normalized_email).first()

def get_user_by_username(db: Session, username: str) -> User | None:
    normalized_username = username.strip().lower()
    return db.query(User).filter(User.username == normalized_username).first()

def validate_password_strength(password: str) -> None:
    if not password or len(password.strip()) < 8:
        raise HTTPException(
            status_code=400,
            detail="Password must be at least 8 characters long",
        )

def require_password_reset_access(
    db: Session,
    *,
    actor: CurrentUser,
    target_user: User,
) -> None:
    if actor.userId == target_user.userId:
        raise HTTPException(status_code=400, detail="Use change password for your own account")

    actor_org_level = normalize_permission_level(actor.organisationPermissionLevel)

    target_org_membership = get_organisation_membership(
        db,
        user_id=target_user.userId,
        organisation_id=actor.organisationId,
    )
    target_org_level = normalize_permission_level(
        target_org_membership.permissionLevel if target_org_membership else None
    )

    if actor_org_level == "admin":
        return

    selected_clinic_id = require_selected_clinic(actor)
    actor_effective_level = get_effective_permission_level(
        db,
        user=actor,
        clinic_id=selected_clinic_id,
    )

    if not actor_effective_level or not has_level(actor_effective_level, "manage"):
        raise HTTPException(status_code=403, detail="You do not have permission to reset passwords")

    if target_org_level == "admin":
        raise HTTPException(status_code=403, detail="You do not have permission to reset this user's password")

    target_effective_level = get_effective_permission_level_for_target_user(
        db,
        target_user=target_user,
        organisation_id=actor.organisationId,
        clinic_id=selected_clinic_id,
    )

    if not target_effective_level:
        raise HTTPException(status_code=403, detail="You do not have permission to reset this user's password")

    if target_effective_level not in {"manage", "write", "read"}:
        raise HTTPException(status_code=403, detail="You do not have permission to reset this user's password")

def get_effective_permission_level_for_target_user(
    db: Session,
    *,
    target_user: User,
    organisation_id: str,
    clinic_id: str | None,
) -> str | None:
    target_org_membership = get_organisation_membership(
        db,
        user_id=target_user.userId,
        organisation_id=organisation_id,
    )
    target_org_level = normalize_permission_level(
        target_org_membership.permissionLevel if target_org_membership else None
    )

    if target_org_level == "admin":
        return "admin"

    if clinic_id:
        target_clinic_membership = get_clinic_membership(
            db,
            user_id=target_user.userId,
            organisation_id=organisation_id,
            clinic_id=clinic_id,
        )
        target_clinic_level = normalize_permission_level(
            target_clinic_membership.permissionLevel if target_clinic_membership else None
        )
        if target_clinic_level:
            return target_clinic_level

    return target_org_level

def organisation_has_admin(db: Session, organisation_id: str) -> bool:
    existing_admin = (
        db.query(OrganisationMembership)
        .filter(
            OrganisationMembership.organisationId == organisation_id,
            OrganisationMembership.permissionLevel == "admin",
        )
        .first()
    )
    return existing_admin is not None

def count_active_organisation_admins(db: Session, organisation_id: str) -> int:
    return (
        db.query(OrganisationMembership)
        .join(User, User.userId == OrganisationMembership.userId)
        .filter(
            OrganisationMembership.organisationId == organisation_id,
            OrganisationMembership.permissionLevel == "admin",
            User.status == "active",
        )
        .count()
    )

def is_last_active_organisation_admin(
    db: Session,
    *,
    organisation_id: str,
    target_user: User,
) -> bool:
    org_membership = get_organisation_membership(
        db,
        user_id=target_user.userId,
        organisation_id=organisation_id,
    )
    target_org_level = normalize_permission_level(
        org_membership.permissionLevel if org_membership else None
    )

    if target_org_level != "admin":
        return False

    if target_user.status != "active":
        return False

    return count_active_organisation_admins(db, organisation_id) <= 1

def get_visible_users_for_actor(
    db: Session,
    *,
    actor: CurrentUser,
) -> list[User]:
    actor_org_level = normalize_permission_level(actor.organisationPermissionLevel)

    if actor_org_level == "admin":
        org_user_ids = (
            db.query(OrganisationMembership.userId)
            .filter(OrganisationMembership.organisationId == actor.organisationId)
            .all()
        )
        clinic_user_ids = (
            db.query(ClinicMembership.userId)
            .filter(ClinicMembership.organisationId == actor.organisationId)
            .all()
        )

        user_ids = {row[0] for row in org_user_ids} | {row[0] for row in clinic_user_ids}
        if not user_ids:
            return []

        return (
            db.query(User)
            .filter(User.userId.in_(user_ids))
            .order_by(User.displayName.asc())
            .all()
        )

    selected_clinic_id = require_selected_clinic(actor)
    actor_effective_level = get_effective_permission_level(
        db,
        user=actor,
        clinic_id=selected_clinic_id,
    )

    if not actor_effective_level or not has_level(actor_effective_level, "manage"):
        raise HTTPException(status_code=403, detail="You do not have permission to view users")

    clinic_user_ids = (
        db.query(ClinicMembership.userId)
        .filter(
            ClinicMembership.organisationId == actor.organisationId,
            ClinicMembership.clinicId == selected_clinic_id,
        )
        .all()
    )

    user_ids = {row[0] for row in clinic_user_ids}
    if not user_ids:
        return []

    return (
        db.query(User)
        .filter(User.userId.in_(user_ids))
        .order_by(User.displayName.asc())
        .all()
    )

def require_user_status_management_access(
    db: Session,
    *,
    actor: CurrentUser,
    target_user: User,
    next_status: str,
) -> None:
    actor_org_level = normalize_permission_level(actor.organisationPermissionLevel)

    if actor.userId == target_user.userId:
        raise HTTPException(status_code=400, detail="You cannot change your own status")

    target_org_membership = get_organisation_membership(
        db,
        user_id=target_user.userId,
        organisation_id=actor.organisationId,
    )
    target_org_level = normalize_permission_level(
        target_org_membership.permissionLevel if target_org_membership else None
    )

    if actor_org_level == "admin":
        if (
            next_status == "disabled"
            and is_last_active_organisation_admin(
                db,
                organisation_id=actor.organisationId,
                target_user=target_user,
            )
        ):
            raise HTTPException(
                status_code=400,
                detail="At least one active organisation admin must remain assigned",
            )
        return

    selected_clinic_id = require_selected_clinic(actor)
    actor_effective_level = get_effective_permission_level(
        db,
        user=actor,
        clinic_id=selected_clinic_id,
    )

    if not actor_effective_level or not has_level(actor_effective_level, "manage"):
        raise HTTPException(status_code=403, detail="You do not have permission to manage users")

    if target_org_level == "admin":
        raise HTTPException(status_code=403, detail="You cannot change the status of an organisation admin")

    target_clinic_memberships = get_clinic_memberships_for_user(
        db,
        user_id=target_user.userId,
        organisation_id=actor.organisationId,
    )

    target_has_selected_clinic = any(
        membership.clinicId == selected_clinic_id for membership in target_clinic_memberships
    )

    if not target_has_selected_clinic:
        raise HTTPException(status_code=403, detail="You can only manage users in your selected clinic")

def require_user_access_update_management(
    db: Session,
    *,
    actor: CurrentUser,
    target_user: User,
    scope: str,
    permission_level: str | None,
    clinic_access: list[ClinicAccessAssignmentRequest] | None,
) -> tuple[str, str | None, dict[str, str]]:
    if actor.userId == target_user.userId:
        raise HTTPException(status_code=400, detail="You cannot change your own access")

    target_org_membership = get_organisation_membership(
        db,
        user_id=target_user.userId,
        organisation_id=actor.organisationId,
    )
    target_org_level = normalize_permission_level(
        target_org_membership.permissionLevel if target_org_membership else None
    )

    actor_org_level = normalize_permission_level(actor.organisationPermissionLevel)
    normalized_scope = (scope or "").strip().lower()

    if normalized_scope not in {"organisation", "clinic"}:
        raise HTTPException(status_code=400, detail="Invalid scope")

    if actor_org_level == "admin":
        if normalized_scope == "organisation":
            normalized_permission_level = normalize_permission_level(permission_level)

            if not normalized_permission_level:
                raise HTTPException(status_code=400, detail="Invalid permission level")

            if not can_grant_level(actor_org_level, normalized_permission_level):
                raise HTTPException(status_code=403, detail="You cannot assign this permission level")

            if (
                target_org_level == "admin"
                and normalized_permission_level != "admin"
                and is_last_active_organisation_admin(
                    db,
                    organisation_id=actor.organisationId,
                    target_user=target_user,
                )
            ):
                raise HTTPException(
                    status_code=400,
                    detail="At least one active organisation admin must remain assigned",
                )

            return normalized_scope, normalized_permission_level, {}

        clinic_access_map = build_clinic_access_map(clinic_access)

        if not clinic_access_map:
            raise HTTPException(
                status_code=400,
                detail="At least one clinic access assignment is required for clinic scope",
            )

        if (
            target_org_level == "admin"
            and is_last_active_organisation_admin(
                db,
                organisation_id=actor.organisationId,
                target_user=target_user,
            )
        ):
            raise HTTPException(
                status_code=400,
                detail="At least one active organisation admin must remain assigned",
            )

        for clinic_id, clinic_permission_level in clinic_access_map.items():
            if not can_grant_level(actor_org_level, clinic_permission_level):
                raise HTTPException(
                    status_code=403,
                    detail=f"You cannot assign {clinic_permission_level} access for clinic {clinic_id}",
                )

        return normalized_scope, None, clinic_access_map

    selected_clinic_id = require_selected_clinic(actor)
    actor_effective_level = get_effective_permission_level(
        db,
        user=actor,
        clinic_id=selected_clinic_id,
    )

    if not actor_effective_level or not has_level(actor_effective_level, "manage"):
        raise HTTPException(status_code=403, detail="You do not have permission to update user access")

    if normalized_scope != "clinic":
        raise HTTPException(
            status_code=403,
            detail="Only organisation admins can assign organisation-level access",
        )

    if target_org_level == "admin":
        raise HTTPException(status_code=403, detail="You cannot change the access of an organisation admin")

    clinic_access_map = build_clinic_access_map(clinic_access)

    if not clinic_access_map:
        raise HTTPException(
            status_code=400,
            detail="At least one clinic access assignment is required for clinic scope",
        )

    for clinic_id, clinic_permission_level in clinic_access_map.items():
        if not can_actor_manage_clinic(db, actor=actor, clinic_id=clinic_id):
            raise HTTPException(
                status_code=403,
                detail=f"You do not have permission to manage clinic {clinic_id}",
            )

        if not can_grant_level(actor_effective_level, clinic_permission_level):
            raise HTTPException(
                status_code=403,
                detail=f"You cannot assign {clinic_permission_level} access for clinic {clinic_id}",
            )

    return normalized_scope, None, clinic_access_map

def require_user_delete_management_access(
    db: Session,
    *,
    actor: CurrentUser,
    target_user: User,
) -> None:
    if actor.userId == target_user.userId:
        raise HTTPException(status_code=400, detail="You cannot delete your own account")

    target_org_membership = get_organisation_membership(
        db,
        user_id=target_user.userId,
        organisation_id=actor.organisationId,
    )
    target_org_level = normalize_permission_level(
        target_org_membership.permissionLevel if target_org_membership else None
    )

    actor_org_level = normalize_permission_level(actor.organisationPermissionLevel)

    if actor_org_level == "admin":
        if is_last_active_organisation_admin(
            db,
            organisation_id=actor.organisationId,
            target_user=target_user,
        ):
            raise HTTPException(
                status_code=400,
                detail="At least one active organisation admin must remain assigned",
            )
        return

    selected_clinic_id = require_selected_clinic(actor)
    actor_effective_level = get_effective_permission_level(
        db,
        user=actor,
        clinic_id=selected_clinic_id,
    )

    if not actor_effective_level or not has_level(actor_effective_level, "manage"):
        raise HTTPException(status_code=403, detail="You do not have permission to delete users")

    if target_org_level == "admin":
        raise HTTPException(status_code=403, detail="You cannot delete an organisation admin")

    target_clinic_memberships = get_clinic_memberships_for_user(
        db,
        user_id=target_user.userId,
        organisation_id=actor.organisationId,
    )

    target_has_selected_clinic = any(
        membership.clinicId == selected_clinic_id for membership in target_clinic_memberships
    )

    if not target_has_selected_clinic:
        raise HTTPException(status_code=403, detail="You can only delete users in your selected clinic")

def can_actor_manage_clinic(
    db: Session,
    *,
    actor: CurrentUser,
    clinic_id: str,
) -> bool:
    effective_level = get_effective_permission_level(
        db,
        user=actor,
        clinic_id=clinic_id,
    )
    return bool(effective_level and has_level(effective_level, "manage"))

def normalize_clinic_access_list(
    clinic_access: list[ClinicAccessInput] | None,
) -> list[dict]:
    normalized_items: list[dict] = []
    seen_clinic_ids: set[str] = set()

    for item in clinic_access or []:
        clinic_id = str(item.clinicId or "").strip()
        permission_level = normalize_permission_level(item.permissionLevel)

        if not clinic_id:
            raise HTTPException(status_code=400, detail="Clinic ID is required")

        if not permission_level:
            raise HTTPException(status_code=400, detail="Invalid clinic permission level")

        if clinic_id in seen_clinic_ids:
            raise HTTPException(
                status_code=400,
                detail=f"Duplicate clinic access entry for clinic {clinic_id}",
            )

        seen_clinic_ids.add(clinic_id)
        normalized_items.append(
            {
                "clinicId": clinic_id,
                "permissionLevel": permission_level,
            }
        )

    return normalized_items

def validate_clinic_access_assignments_for_create(
    db: Session,
    *,
    actor: CurrentUser,
    clinic_access: list[dict],
) -> list[dict]:
    if not clinic_access:
        raise HTTPException(
            status_code=400,
            detail="At least one clinic access assignment is required for clinic scope",
        )

    actor_org_level = normalize_permission_level(actor.organisationPermissionLevel)

    for item in clinic_access:
        clinic_id = item["clinicId"]
        permission_level = item["permissionLevel"]

        if actor_org_level == "admin":
            if not can_grant_level(actor_org_level, permission_level):
                raise HTTPException(
                    status_code=403,
                    detail=f"You cannot assign {permission_level} for clinic {clinic_id}",
                )
        else:
            actor_effective_level = get_effective_permission_level(
                db,
                user=actor,
                clinic_id=clinic_id,
            )

            if not actor_effective_level or not can_create_users(actor_effective_level):
                raise HTTPException(
                    status_code=403,
                    detail=f"You do not have permission to create users for clinic {clinic_id}",
                )

            if not can_grant_level(actor_effective_level, permission_level):
                raise HTTPException(
                    status_code=403,
                    detail=f"You cannot assign {permission_level} for clinic {clinic_id}",
                )

            if not can_actor_manage_clinic(db, actor=actor, clinic_id=clinic_id):
                raise HTTPException(
                    status_code=403,
                    detail=f"You do not have permission to manage clinic {clinic_id}",
                )

    return clinic_access

def validate_clinic_access_assignments_for_update(
    db: Session,
    *,
    actor: CurrentUser,
    target_user: User,
    clinic_access: list[dict],
) -> list[dict]:
    if not clinic_access:
        raise HTTPException(
            status_code=400,
            detail="At least one clinic access assignment is required for clinic scope",
        )

    target_org_membership = get_organisation_membership(
        db,
        user_id=target_user.userId,
        organisation_id=actor.organisationId,
    )
    target_org_level = normalize_permission_level(
        target_org_membership.permissionLevel if target_org_membership else None
    )

    actor_org_level = normalize_permission_level(actor.organisationPermissionLevel)

    if target_org_level == "admin":
        raise HTTPException(
            status_code=403,
            detail="You cannot change the access of an organisation admin",
        )

    for item in clinic_access:
        clinic_id = item["clinicId"]
        permission_level = item["permissionLevel"]

        if actor_org_level == "admin":
            if not can_grant_level(actor_org_level, permission_level):
                raise HTTPException(
                    status_code=403,
                    detail=f"You cannot assign {permission_level} for clinic {clinic_id}",
                )
        else:
            actor_effective_level = get_effective_permission_level(
                db,
                user=actor,
                clinic_id=clinic_id,
            )

            if not actor_effective_level or not has_level(actor_effective_level, "manage"):
                raise HTTPException(
                    status_code=403,
                    detail=f"You do not have permission to update access for clinic {clinic_id}",
                )

            if not can_grant_level(actor_effective_level, permission_level):
                raise HTTPException(
                    status_code=403,
                    detail=f"You cannot assign {permission_level} for clinic {clinic_id}",
                )

            if not can_actor_manage_clinic(db, actor=actor, clinic_id=clinic_id):
                raise HTTPException(
                    status_code=403,
                    detail=f"You do not have permission to manage clinic {clinic_id}",
                )

    return clinic_access

def normalize_clinic_access_assignments(
    clinic_access: list[ClinicAccessAssignmentRequest] | None,
) -> list[tuple[str, str]]:
    normalized_items: list[tuple[str, str]] = []
    seen_clinic_ids: set[str] = set()

    for item in clinic_access or []:
        clinic_id = (item.clinicId or "").strip()
        permission_level = normalize_permission_level(item.permissionLevel)

        if not clinic_id:
            raise HTTPException(status_code=400, detail="Clinic ID is required")

        if not permission_level:
            raise HTTPException(status_code=400, detail="Invalid permission level")

        if clinic_id in seen_clinic_ids:
            raise HTTPException(
                status_code=400,
                detail=f"Clinic {clinic_id} was added more than once",
            )

        seen_clinic_ids.add(clinic_id)
        normalized_items.append((clinic_id, permission_level))

    return normalized_items

def build_clinic_access_map(
    clinic_access: list[ClinicAccessAssignmentRequest] | None,
) -> dict[str, str]:
    normalized_items = normalize_clinic_access_assignments(clinic_access)
    return {clinic_id: permission_level for clinic_id, permission_level in normalized_items}

def require_user_management_access(
    db: Session,
    *,
    actor: CurrentUser,
    scope: str,
    permission_level: str | None,
    clinic_access: list[ClinicAccessAssignmentRequest] | None,
) -> tuple[str, str | None, dict[str, str]]:
    normalized_scope = (scope or "").strip().lower()
    actor_org_level = normalize_permission_level(actor.organisationPermissionLevel)

    if normalized_scope not in {"organisation", "clinic"}:
        raise HTTPException(status_code=400, detail="Invalid scope")

    if normalized_scope == "organisation":
        normalized_permission_level = normalize_permission_level(permission_level)

        if not normalized_permission_level:
            raise HTTPException(status_code=400, detail="Invalid permission level")

        if actor_org_level != "admin":
            raise HTTPException(
                status_code=403,
                detail="Only organisation admins can assign organisation-level access",
            )

        if not can_grant_level(actor_org_level, normalized_permission_level):
            raise HTTPException(status_code=403, detail="You cannot assign this permission level")

        return normalized_scope, normalized_permission_level, {}

    clinic_access_map = build_clinic_access_map(clinic_access)

    if not clinic_access_map:
        raise HTTPException(
            status_code=400,
            detail="At least one clinic access assignment is required for clinic scope",
        )

    selected_clinic_id = require_selected_clinic(actor)
    actor_effective_level = get_effective_permission_level(
        db,
        user=actor,
        clinic_id=selected_clinic_id,
    )

    if not actor_effective_level or not can_create_users(actor_effective_level):
        raise HTTPException(status_code=403, detail="You do not have permission to create users")

    for clinic_id, clinic_permission_level in clinic_access_map.items():
        if not can_actor_manage_clinic(
            db,
            actor=actor,
            clinic_id=clinic_id,
        ):
            raise HTTPException(
                status_code=403,
                detail=f"You do not have permission to manage clinic {clinic_id}",
            )

        if not can_grant_level(actor_effective_level, clinic_permission_level):
            raise HTTPException(
                status_code=403,
                detail=f"You cannot assign {clinic_permission_level} access for clinic {clinic_id}",
            )

    return normalized_scope, None, clinic_access_map

def get_effective_permission_level(
    db: Session,
    *,
    user: CurrentUser,
    clinic_id: str | None,
) -> str | None:
    require_organisation_access(user, user.organisationId)

    org_membership = get_organisation_membership(
        db,
        user_id=user.userId,
        organisation_id=user.organisationId,
    )

    org_level = normalize_permission_level(
        org_membership.permissionLevel if org_membership else user.organisationPermissionLevel
    )

    if org_level == "admin":
        return "admin"

    if clinic_id:
        clinic_membership = get_clinic_membership(
            db,
            user_id=user.userId,
            organisation_id=user.organisationId,
            clinic_id=clinic_id,
        )
        clinic_level = normalize_permission_level(
            clinic_membership.permissionLevel if clinic_membership else None
        )

        if clinic_level:
            return clinic_level

    return org_level

def require_effective_permission(
    db: Session,
    *,
    user: CurrentUser,
    clinic_id: str | None,
    required_level: str,
) -> str:
    effective_level = get_effective_permission_level(
        db,
        user=user,
        clinic_id=clinic_id,
    )

    if not effective_level or not has_level(effective_level, required_level):
        raise HTTPException(status_code=403, detail="You do not have permission to perform this action")

    return effective_level

def require_document_access(
    db: Session,
    *,
    user: CurrentUser,
    document: Document,
    required_level: str = "read",
) -> str:
    if document.organisationId != user.organisationId:
        raise HTTPException(status_code=403, detail="You do not have access to this document")

    effective_level = get_effective_permission_level(
        db,
        user=user,
        clinic_id=document.clinicId,
    )

    if not effective_level:
        raise HTTPException(status_code=403, detail="You do not have access to this document")

    if not has_level(effective_level, required_level):
        raise HTTPException(status_code=403, detail="You do not have permission to perform this action")

    if not can_view_document(effective_level, document.roleAccess):
        raise HTTPException(status_code=403, detail="You do not have access to this document")

    if document.isShared and not document.clinicId:
        if document.organisationId != user.organisationId:
            raise HTTPException(status_code=403, detail="You do not have access to this document")

    if document.clinicId:
        clinic_membership = get_clinic_membership(
            db,
            user_id=user.userId,
            organisation_id=user.organisationId,
            clinic_id=document.clinicId,
        )
        org_membership = get_organisation_membership(
            db,
            user_id=user.userId,
            organisation_id=user.organisationId,
        )

        has_org_admin = normalize_permission_level(
            org_membership.permissionLevel if org_membership else user.organisationPermissionLevel
        ) == "admin"

        if not clinic_membership and not has_org_admin:
            raise HTTPException(status_code=403, detail="You do not have access to this clinic document")

    return effective_level

def get_allowed_document_levels_for_user(
    db: Session,
    *,
    user: CurrentUser,
    clinic_id: str | None,
) -> list[str]:
    effective_level = get_effective_permission_level(
        db,
        user=user,
        clinic_id=clinic_id,
    )

    return get_accessible_document_levels(effective_level)

def can_user_access_shared_scope(
    db: Session,
    *,
    user: CurrentUser,
    clinic_id: str | None,
) -> bool:
    effective_level = get_effective_permission_level(
        db,
        user=user,
        clinic_id=clinic_id,
    )

    return bool(effective_level and can_view_document(effective_level, "read"))

STORAGE_PROVIDER = os.getenv("STORAGE_PROVIDER", "local").strip().lower()
LOCAL_UPLOAD_ROOT = Path(
    os.getenv("LOCAL_UPLOAD_ROOT", r"C:\dev\dba-mvp\data\uploads")
).expanduser()

QDRANT_URL = os.getenv("QDRANT_URL", "http://localhost:6333")
COLLECTION = os.getenv("QDRANT_COLLECTION", "dba_docs")

app = FastAPI(title="DBA MVP API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

app.include_router(chat_router)
app.include_router(insights_router)

# -------- Qdrant setup --------
qdrant = QdrantClient(url=QDRANT_URL)
VECTOR_SIZE = 1536

def ensure_collection():
    existing = [c.name for c in qdrant.get_collections().collections]
    if COLLECTION not in existing:
        qdrant.create_collection(
            collection_name=COLLECTION,
            vectors_config=qm.VectorParams(
                size=VECTOR_SIZE,
                distance=qm.Distance.COSINE,
            ),
        )

ensure_collection()

# -------- Helpers --------
def build_storage_path(
    *,
    organisation_id: str,
    clinic_id: str | None,
    filename: str,
) -> str:
    folder = clinic_id if clinic_id else "shared"
    safe_filename = Path(filename).name
    return str(Path(organisation_id) / folder / safe_filename)

def get_prompt_clinic_name(db: Session, current_user: CurrentUser) -> str:
    if current_user.selectedClinicId:
        clinic = (
            db.query(Clinic)
            .filter(
                Clinic.clinicId == current_user.selectedClinicId,
                Clinic.organisationId == current_user.organisationId,
            )
            .first()
        )
        if clinic and clinic.name:
            return clinic.name

    return "the clinic"

def build_base_prompt(clinic_name: str) -> str:
    return f"""
You are Dental Buddy AI, an AI assistant for {clinic_name}.

Be clear, calm, natural, and helpful.
Keep answers concise and easy to follow.
Do not sound robotic or overly formal.
Always preserve the clinic name exactly as written: {clinic_name}
""".strip()

def build_identity_prompt(clinic_name: str) -> str:
    base_prompt = build_base_prompt(clinic_name)
    return f"""
{base_prompt}

The user is asking who you are or what you do.

Answer naturally and briefly.

Important:
- Do not assume capabilities like booking, treatments, diagnosis, or clinical actions
- Only describe your actual role
- Do not overstate what you can do

Your role is to help answer questions using the clinic's documents.

Do not over-explain.
Do not sound like marketing.
""".strip()

def build_greeting_prompt(clinic_name: str) -> str:
    base_prompt = build_base_prompt(clinic_name)
    return f"""
{base_prompt}

The user has greeted you.

Reply naturally and briefly.
Be warm but not overly friendly.
Then gently ask if they need help.

Do not mention documents unless asked.
""".strip()

def build_help_prompt(clinic_name: str) -> str:
    base_prompt = build_base_prompt(clinic_name)
    return f"""
{base_prompt}

The user is asking what you can do or how you can help.

Answer naturally and briefly.

Important:
- Do not assume capabilities like booking, treatments, diagnosis, or clinical actions
- Only describe what you actually do
- Do not claim to perform tasks outside answering questions from clinic documents

You help by:
- answering questions using the clinic's documents
- helping staff find procedures, policies, forms, and internal guidance

Keep it simple and accurate.
Do not over-explain.
Do not sound like marketing.
""".strip()

def build_rag_prompt(clinic_name: str) -> str:
    base_prompt = build_base_prompt(clinic_name)
    return f"""
{base_prompt}

You are answering a user question using the clinic documents provided.

Document rules:
- Use the provided clinic documents as the source of truth
- Do not invent clinic rules, procedures, or policies
- If the answer is not clearly in the documents, say so
- If the documents only partly answer the question, say what is clear and what is missing

How to answer:
- Answer the user’s exact question
- Keep the answer practical, short, and easy to follow
- Prefer a short opening sentence first
- Group related items under small headings when there are multiple items
- Use bullet points for lists
- Keep each bullet short
- Do not dump long flat lists without grouping
- Do not write long paragraphs
- If the answer is a list of tools, platforms, steps, or roles, organise them into logical categories
- If there are more than 6 list items, group them into categories instead of showing one long list

If the clinic documents do not contain the answer:
- Say that clearly
- Tell the user to contact their Practice Manager or Clinic Owner
""".strip()

def get_local_absolute_path(storage_path: str) -> Path:
    return LOCAL_UPLOAD_ROOT / Path(storage_path)

def save_file_to_storage(
    *,
    organisation_id: str,
    clinic_id: str | None,
    filename: str,
    contents: bytes,
) -> tuple[str, str, Path]:
    if STORAGE_PROVIDER != "local":
        raise RuntimeError(f"Unsupported STORAGE_PROVIDER for now: {STORAGE_PROVIDER}")

    storage_path = build_storage_path(
        organisation_id=organisation_id,
        clinic_id=clinic_id,
        filename=filename,
    )
    absolute_path = get_local_absolute_path(storage_path)
    absolute_path.parent.mkdir(parents=True, exist_ok=True)

    with open(absolute_path, "wb") as f:
        f.write(contents)

    return "local", storage_path, absolute_path

def get_document_absolute_path(document: Document) -> Path:
    if document.storageProvider and document.storagePath:
        if document.storageProvider != "local":
            raise RuntimeError(f"Unsupported storage provider for local file access: {document.storageProvider}")
        return get_local_absolute_path(document.storagePath)

    # fallback for old records created before storagePath existed
    clinic_folder = document.clinicId if document.clinicId else "shared"
    return LOCAL_UPLOAD_ROOT / document.organisationId / clinic_folder / document.filename

def delete_document_from_storage(document: Document) -> None:
    path = get_document_absolute_path(document)
    if path.exists():
        path.unlink()

def build_document_qdrant_filter(document_id: str) -> qm.Filter:
    return qm.Filter(
        must=[
            qm.FieldCondition(
                key="documentId",
                match=qm.MatchValue(value=document_id),
            )
        ]
    )

def delete_document_vectors(document_id: str) -> None:
    qdrant.delete(
        collection_name=COLLECTION,
        points_selector=qm.FilterSelector(
            filter=build_document_qdrant_filter(document_id),
        ),
    )

def set_document_vectors_status(document_id: str, status: str) -> None:
    qdrant.set_payload(
        collection_name=COLLECTION,
        payload={"status": status},
        points=build_document_qdrant_filter(document_id),
    )

def count_document_vectors(document_id: str) -> int:
    response = qdrant.count(
        collection_name=COLLECTION,
        count_filter=build_document_qdrant_filter(document_id),
        exact=True,
    )
    return int(response.count or 0)

def evaluate_document_consistency(document: Document) -> dict:
    vector_count = count_document_vectors(document.documentId)
    issues: list[str] = []

    status = (document.status or "").strip().lower()
    index_status = (document.indexStatus or "").strip().lower()

    if status == "active":
        if index_status == "indexed" and vector_count == 0:
            issues.append("Document is indexed in DB but has no vectors in Qdrant")

        if index_status == "pending" and vector_count > 0:
            issues.append("Document is still pending in DB but already has vectors in Qdrant")

        if index_status == "failed" and vector_count > 0:
            issues.append("Document is marked failed in DB but still has vectors in Qdrant")

        if index_status == "stale" and vector_count == 0:
            issues.append("Document is marked stale in DB and has no vectors in Qdrant")

    if status == "archived" and vector_count > 0:
        issues.append("Document is archived in DB but still has vectors in Qdrant")

    if status not in {"active", "archived"} and vector_count > 0:
        issues.append("Document has vectors in Qdrant but DB status is not active/archived")

    return {
        "documentId": document.documentId,
        "filename": document.filename,
        "clinicId": document.clinicId,
        "isShared": bool(document.isShared),
        "status": document.status,
        "indexStatus": document.indexStatus,
        "vectorCount": vector_count,
        "isConsistent": len(issues) == 0,
        "issues": issues,
    }

def repair_document_by_id(document_id: str) -> dict:
    db = SessionLocal()
    document = None

    try:
        document = db.query(Document).filter(Document.documentId == document_id).first()

        if not document:
            return {
                "status": "error",
                "message": "Document not found in database",
            }

        status = (document.status or "").strip().lower()

        if status == "archived":
            delete_document_vectors(document.documentId)
            document.indexStatus = "stale"
            document.indexError = None
            db.commit()

            return {
                "status": "repaired",
                "action": "deleted_vectors_for_archived_document",
                "documentId": document.documentId,
                "filename": document.filename,
                "vectorCount": count_document_vectors(document.documentId),
            }

        if status == "active":
            delete_document_vectors(document.documentId)
            document.indexStatus = "pending"
            document.indexError = None
            db.commit()

            result = process_document_by_id(document.documentId)

            return {
                "status": "repaired" if result.get("status") == "ingested" else "error",
                "action": "reindexed_active_document",
                "documentId": document.documentId,
                "filename": document.filename,
                "reindexResult": result,
                "vectorCount": count_document_vectors(document.documentId),
            }

        delete_document_vectors(document.documentId)
        document.indexStatus = "stale"
        document.indexError = f"Repair removed vectors for unsupported document status: {document.status}"
        db.commit()

        return {
            "status": "repaired",
            "action": "deleted_vectors_for_non_active_document",
            "documentId": document.documentId,
            "filename": document.filename,
            "vectorCount": count_document_vectors(document.documentId),
        }

    except Exception as e:
        if document:
            document.indexStatus = "failed"
            document.indexError = str(e)[:1000]
            db.commit()

        return {
            "status": "error",
            "message": str(e),
        }
    finally:
        db.close()

def get_documents_in_scope_for_current_user(
    db: Session,
    *,
    current_user: CurrentUser,
) -> list[Document]:
    selected_clinic_id = current_user.selectedClinicId

    query = db.query(Document).filter(
        Document.organisationId == current_user.organisationId,
    )

    org_membership = get_organisation_membership(
        db,
        user_id=current_user.userId,
        organisation_id=current_user.organisationId,
    )
    org_level = normalize_permission_level(
        org_membership.permissionLevel if org_membership else current_user.organisationPermissionLevel
    )
    is_org_admin = org_level == "admin"

    if is_org_admin:
        if selected_clinic_id:
            query = query.filter(
                (Document.clinicId == selected_clinic_id)
                | (Document.isShared == True)
            )
    else:
        selected_clinic_id = require_selected_clinic(current_user)

        clinic_membership = get_clinic_membership(
            db,
            user_id=current_user.userId,
            organisation_id=current_user.organisationId,
            clinic_id=selected_clinic_id,
        )
        if not clinic_membership:
            raise HTTPException(status_code=403, detail="You do not have access to this clinic")

        query = query.filter(
            (Document.clinicId == selected_clinic_id)
            | (Document.isShared == True)
        )

    return query.order_by(Document.uploadedAt.desc()).all()

def get_orphaned_vector_document_ids_for_org(organisation_id: str) -> list[str]:
    orphan_counts: dict[str, int] = {}
    offset = None

    while True:
        points, next_offset = qdrant.scroll(
            collection_name=COLLECTION,
            scroll_filter=qm.Filter(
                must=[
                    qm.FieldCondition(
                        key="organisationId",
                        match=qm.MatchValue(value=organisation_id),
                    )
                ]
            ),
            with_payload=True,
            with_vectors=False,
            limit=200,
            offset=offset,
        )

        if not points:
            break

        for point in points:
            payload = point.payload or {}
            document_id = payload.get("documentId")
            if not document_id:
                continue
            orphan_counts[document_id] = orphan_counts.get(document_id, 0) + 1

        if next_offset is None:
            break

        offset = next_offset

    return list(orphan_counts.keys())

def get_orphaned_vectors_for_org(db: Session, organisation_id: str) -> list[dict]:
    db_document_ids = {
        row[0]
        for row in db.query(Document.documentId)
        .filter(Document.organisationId == organisation_id)
        .all()
    }

    qdrant_document_ids = get_orphaned_vector_document_ids_for_org(organisation_id)

    orphaned = []
    for document_id in qdrant_document_ids:
        if document_id not in db_document_ids:
            orphaned.append(
                {
                    "documentId": document_id,
                    "vectorCount": count_document_vectors(document_id),
                }
            )

    orphaned.sort(key=lambda x: x["vectorCount"], reverse=True)
    return orphaned

def build_document_health_summary(
    db: Session,
    *,
    current_user: CurrentUser,
) -> dict:
    documents = get_documents_in_scope_for_current_user(
        db,
        current_user=current_user,
    )

    total_documents = len(documents)
    active_documents = 0
    archived_documents = 0
    indexed_documents = 0
    pending_documents = 0
    failed_documents = 0
    stale_documents = 0
    inconsistent_documents = 0

    for document in documents:
        status = (document.status or "").strip().lower()
        index_status = (document.indexStatus or "").strip().lower()

        if status == "active":
            active_documents += 1
        elif status == "archived":
            archived_documents += 1

        if index_status == "indexed":
            indexed_documents += 1
        elif index_status == "pending":
            pending_documents += 1
        elif index_status == "failed":
            failed_documents += 1
        elif index_status == "stale":
            stale_documents += 1

        consistency = evaluate_document_consistency(document)
        if not consistency["isConsistent"]:
            inconsistent_documents += 1

    orphaned_vector_groups = len(
        get_orphaned_vectors_for_org(db, current_user.organisationId)
    )

    return {
        "totalDocuments": total_documents,
        "activeDocuments": active_documents,
        "archivedDocuments": archived_documents,
        "indexedDocuments": indexed_documents,
        "pendingDocuments": pending_documents,
        "failedDocuments": failed_documents,
        "staleDocuments": stale_documents,
        "inconsistentDocuments": inconsistent_documents,
        "orphanedVectorGroups": orphaned_vector_groups,
    }

def reindex_documents_in_scope(
    db: Session,
    *,
    current_user: CurrentUser,
    only_non_indexed: bool,
    limit: int,
) -> dict:
    documents = get_documents_in_scope_for_current_user(
        db,
        current_user=current_user,
    )

    safe_limit = max(1, min(limit, 200))

    checked = 0
    reindexed = 0
    skipped = 0
    results = []

    for document in documents:
        if checked >= safe_limit:
            break

        checked += 1

        status = (document.status or "").strip().lower()
        index_status = (document.indexStatus or "").strip().lower()

        if status != "active":
            skipped += 1
            continue

        if only_non_indexed and index_status == "indexed":
            skipped += 1
            continue

        try:
            require_document_access(
                db,
                user=current_user,
                document=document,
                required_level="manage",
            )
        except HTTPException:
            skipped += 1
            continue

        repair_result = repair_document_by_id(document.documentId)

        results.append(
            {
                "documentId": document.documentId,
                "filename": document.filename,
                "status": document.status,
                "indexStatus": document.indexStatus,
                "repairResult": repair_result,
            }
        )

        if repair_result.get("status") == "repaired":
            reindexed += 1

    return {
        "status": "completed",
        "checked": checked,
        "reindexed": reindexed,
        "skipped": skipped,
        "limit": safe_limit,
        "onlyNonIndexed": only_non_indexed,
        "results": results,
    }

def normalize_line(text: str) -> str:
    return " ".join(text.strip().split())

def fallback_chunk_text(text: str, chunk_size: int = 1800, overlap: int = 400) -> List[str]:
    text = normalize_line(text)
    chunks = []
    i = 0
    while i < len(text):
        chunk = text[i:i + chunk_size]
        chunks.append(chunk)
        i += max(1, chunk_size - overlap)
    return chunks

def is_step_heading(line: str) -> bool:
    line = normalize_line(line)
    return bool(re.match(r"^step\s+\d+\b", line, flags=re.IGNORECASE))

def extract_step_number_from_heading(line: str) -> Optional[str]:
    match = re.match(r"^step\s+(\d+)\b", normalize_line(line), flags=re.IGNORECASE)
    if match:
        return match.group(1)
    return None

def is_section_heading(line: str) -> bool:
    line = normalize_line(line)
    lower = line.lower()


    known_headings = {
        "overview",
        "goal",
        "purpose",
        "scope",
        "procedure",
        "procedures",
        "reporting",
        "notification",
        "notify",
        "responsibilities",
        "responsibility",
        "key behaviours",
        "key behavior",
        "tools and supports",
        "tools & supports",
        "definitions",
        "references",
        "follow up",
        "follow-up",
        "medical follow-up",
        "documentation",
        "documentation and follow-up",
        "preventive measures",
    }


    if lower in known_headings:
        return True


    # short title-like headings only
    if len(line) <= 50 and line == line.title():
        return True


    return False

def extract_text_from_file(path: Path) -> str:
    ext = path.suffix.lower()


    if ext == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore")


    if ext == ".pdf":
        reader = PdfReader(str(path))
        pages = []
        for p in reader.pages:
            pages.append(p.extract_text() or "")
        return "\n".join(pages)


    if ext == ".docx":
        docx_file = DocxDocument(str(path))
        parts = []


        # paragraphs first
        for para in docx_file.paragraphs:
            text = normalize_line(para.text)
            if text:
                parts.append(text)


        # tables after paragraphs
        for table in docx_file.tables:
            for row in table.rows:
                row_parts = []
                for cell in row.cells:
                    cell_text = normalize_line(cell.text)
                    if cell_text:
                        row_parts.append(cell_text)
                if row_parts:
                    parts.append(" | ".join(row_parts))


        return "\n".join(parts)


    return ""

def build_sop_chunks_from_lines(lines: List[str]) -> List[dict]:
    """
    Returns structured chunks with metadata:
    [
      {
        "title": "Procedure — Step 1",
        "section": "Procedure",
        "step_number": "1",
        "text": "..."
      }
    ]
    """
    chunks: List[dict] = []


    if not lines:
        return chunks


    metadata_lines = []
    body_start = 0


    # everything before first real heading goes to metadata
    for i, line in enumerate(lines):
        if is_section_heading(line) or is_step_heading(line):
            body_start = i
            break
        metadata_lines.append(line)
        body_start = i + 1


    if metadata_lines:
        chunks.append(
            {
                "title": "Metadata",
                "section": "Metadata",
                "step_number": None,
                "text": "\n".join(metadata_lines),
            }
        )


    current_section: Optional[str] = None
    current_step: Optional[str] = None
    current_lines: List[str] = []


    def flush_chunk():
        nonlocal current_section, current_step, current_lines


        if not current_lines and not current_section and not current_step:
            return


        title_parts = []
        if current_section:
            title_parts.append(current_section)
        if current_step:
            title_parts.append(current_step)


        title = " — ".join(title_parts).strip() or "Content"
        body = "\n".join(current_lines).strip()


        if body:
            chunks.append(
                {
                    "title": title,
                    "section": current_section,
                    "step_number": extract_step_number_from_heading(current_step) if current_step else None,
                    "text": body,
                }
            )


        current_step = None
        current_lines = []


    for line in lines[body_start:]:
        if is_step_heading(line):
            flush_chunk()
            current_step = normalize_line(line)
            current_lines = []
            continue


        if is_section_heading(line):
            flush_chunk()
            current_section = normalize_line(line)
            current_step = None
            current_lines = []
            continue


        current_lines.append(line)


    flush_chunk()
    return chunks

def structure_aware_chunk_text(text: str) -> List[dict]:
    lines = [normalize_line(line) for line in text.splitlines() if normalize_line(line)]
    structured_chunks = build_sop_chunks_from_lines(lines)


    # if not enough structure detected, fallback later
    return structured_chunks

def embed_text(text: str) -> List[float]:
    response = azure_client.embeddings.create(
        model=EMBEDDING_DEPLOYMENT,
        input=text,
    )
    return response.data[0].embedding

def rerank_sources(question: str, sources: List[dict], max_sources: int = 4) -> List[dict]:
    if not sources:
        return []


    numbered_sources = "\n\n".join(
        [
            (
                f"[{i+1}] Title: {s.get('title', '')}\n"
                f"Section: {s.get('section', '')}\n"
                f"Step: {s.get('stepNumber', '')}\n"
                f"Filename: {s['filename']}\n"
                f"Text: {s['text']}"
            )
            for i, s in enumerate(sources)
        ]
    )


    response = azure_client.chat.completions.create(
        model=CHAT_DEPLOYMENT,
        messages=[
            {
                "role": "system",
                "content": """
You are helping rank document chunks for a dental clinic AI assistant.


Given a user question and a list of retrieved document chunks, choose the most relevant chunks for answering the question.


Prioritize exact step matches when the user asks for a specific step number.


Return only a comma-separated list of chunk numbers in best-first order.
Example: 3,1,4,2


Do not explain your answer.
"""
            },
            {
                "role": "user",
                "content": f"""
Question:
{question}


Retrieved chunks:
{numbered_sources}


Return the top {max_sources} most relevant chunk numbers only.
"""
            }
        ],
        temperature=0,
    )


    content = (response.choices[0].message.content or "").strip()


    try:
        chosen_indexes = []
        for part in content.split(","):
            n = int(part.strip())
            if 1 <= n <= len(sources):
                chosen_indexes.append(n - 1)


        seen = set()
        unique_indexes = []
        for idx in chosen_indexes:
            if idx not in seen:
                seen.add(idx)
                unique_indexes.append(idx)


        reranked = [sources[idx] for idx in unique_indexes[:max_sources]]
        if reranked:
            return reranked
    except Exception:
        pass


    return sources[:max_sources]

def apply_question_heuristics(question: str, sources: List[dict]) -> List[dict]:
    if not sources:
        return sources


    q = question.lower()


    # exact step preference
    step_match = re.search(r"\bstep\s+(\d+)\b", q)
    if step_match:
        wanted = step_match.group(1)
        matching = [s for s in sources if str(s.get("stepNumber") or "") == wanted]
        non_matching = [s for s in sources if str(s.get("stepNumber") or "") != wanted]
        if matching:
            return matching + non_matching


    # notify / report preference
    notify_words = ["notify", "notification", "report", "reported", "manager", "principal dentist"]
    if any(word in q for word in notify_words):
        matching = [
            s for s in sources
            if any(word in (s.get("text", "") + " " + (s.get("title") or "")).lower() for word in notify_words)
        ]
        non_matching = [s for s in sources if s not in matching]
        if matching:
            return matching + non_matching


    return sources

def assess_document_readiness(text: str, structured_chunks: List[dict]) -> dict:
    text_length = len(text.strip())
    has_step_structure = any(chunk.get("step_number") for chunk in structured_chunks)
    has_sections = any(chunk.get("section") for chunk in structured_chunks if chunk.get("section"))
    has_enough_text = text_length >= 400
    has_multiple_chunks = len(structured_chunks) >= 2

    notes = []

    if has_enough_text:
        notes.append("Readable text extracted")
    else:
        notes.append("Limited readable text extracted")

    if has_sections:
        notes.append("Section headings detected")
    else:
        notes.append("Limited section structure detected")

    if has_step_structure:
        notes.append("Step structure detected")

    if not has_multiple_chunks:
        notes.append("Document structure appears limited")

    if not has_enough_text:
        readiness = "Needs improvement"
    elif has_sections and has_multiple_chunks:
        if has_step_structure:
            readiness = "Good for DBA"
        else:
            readiness = "Usable with warnings"
    else:
        readiness = "Usable with warnings"

    return {
        "readiness": readiness,
        "notes": notes,
    }

def process_document_by_id(document_id: str):
    db = SessionLocal()
    document = None

    try:
        document = db.query(Document).filter(Document.documentId == document_id).first()

        if not document:
            return {"status": "error", "message": "Document not found in database"}

        doc_path = get_document_absolute_path(document)

        if not doc_path.exists():
            document.indexStatus = "failed"
            document.indexError = "File not found on disk"
            db.commit()
            return {
                "status": "error",
                "message": "File not found on disk",
                "path": str(doc_path),
            }

        text = extract_text_from_file(doc_path)
        if not text.strip():
            document.indexStatus = "failed"
            document.indexError = "No extractable text"
            db.commit()
            return {
                "status": "error",
                "message": "No extractable text (try .txt or text-based PDF)",
            }

        structured_chunks = structure_aware_chunk_text(text)
        readiness_result = assess_document_readiness(text, structured_chunks)

        points = []

        if len(structured_chunks) >= 2:
            for idx, chunk in enumerate(structured_chunks):
                title = chunk.get("title") or "Content"
                section = chunk.get("section") or ""
                step_number = chunk.get("step_number")
                body_text = chunk.get("text") or ""

                enriched_text = f"""
Document: {document.filename}
Document Type: {document.documentType}
Chunk Title: {title}
Section: {section}
Step Number: {step_number if step_number else ""}

{body_text}
"""

                vec = embed_text(enriched_text)
                point_id = str(uuid.uuid4())

                payload = {
                    "documentId": document.documentId,
                    "organisationId": document.organisationId,
                    "clinicId": document.clinicId,
                    "filename": document.filename,
                    "documentType": document.documentType,
                    "roleAccess": document.roleAccess,
                    "sourceType": document.sourceType,
                    "isShared": document.isShared,
                    "status": "active",
                    "chunkIndex": idx,
                    "sourceId": f"{document.documentId}:{idx}",
                    "title": title,
                    "section": section,
                    "stepNumber": step_number,
                    "text": body_text,
                }

                points.append(qm.PointStruct(id=point_id, vector=vec, payload=payload))
        else:
            fallback_chunks = fallback_chunk_text(text)
            for idx, chunk in enumerate(fallback_chunks):
                enriched_text = f"""
Document: {document.filename}
Document Type: {document.documentType}

{chunk}
"""
                vec = embed_text(enriched_text)
                point_id = str(uuid.uuid4())

                payload = {
                    "documentId": document.documentId,
                    "organisationId": document.organisationId,
                    "clinicId": document.clinicId,
                    "filename": document.filename,
                    "documentType": document.documentType,
                    "roleAccess": document.roleAccess,
                    "sourceType": document.sourceType,
                    "isShared": document.isShared,
                    "status": "active",
                    "chunkIndex": idx,
                    "sourceId": f"{document.documentId}:{idx}",
                    "title": None,
                    "section": None,
                    "stepNumber": None,
                    "text": chunk,
                }

                points.append(qm.PointStruct(id=point_id, vector=vec, payload=payload))

        qdrant.upsert(collection_name=COLLECTION, points=points)
        document.readiness = readiness_result["readiness"]
        document.readinessNotes = ", ".join(readiness_result["notes"])
        document.status = "active"
        document.indexStatus = "indexed"
        document.indexError = None
        document.indexedAt = datetime.utcnow()
        db.commit()

        return {
            "status": "ingested",
            "documentId": document.documentId,
            "organisationId": document.organisationId,
            "clinicId": document.clinicId,
            "filename": document.filename,
            "chunks": len(points),
            "collection": COLLECTION,
            "readiness": readiness_result["readiness"],
            "readinessNotes": readiness_result["notes"],
        }

    except Exception as e:
        if document:
            document.indexStatus = "failed"
            document.indexError = str(e)[:1000]
            db.commit()

        return {
            "status": "error",
            "message": str(e),
        }
    finally:
        db.close()

def process_document_background(document_id: str):
    result = process_document_by_id(document_id)
    if result.get("status") != "ingested":
        print(f"Background processing failed for {document_id}: {result}")

CLARIFICATION_SUGGESTIONS = {
    "infection": [
        "Needle stick injury",
        "Cross contamination",
        "Patient infection management",
    ],
    "sterilisation": [
        "Sterilisation cycle process",
        "Failed sterilisation test",
        "Instrument cleaning before sterilisation",
    ],
    "emergency": [
        "Medical emergency",
        "After-hours emergency contact",
        "Incident reporting after an emergency",
    ],
    "consent": [
        "Treatment consent",
        "Financial consent",
        "Consent documentation requirements",
    ],
    "complaint": [
        "Patient complaint process",
        "Complaint escalation",
        "Complaint documentation",
    ],
    "incident": [
        "Needle stick injury",
        "Workplace incident reporting",
        "Patient safety incident",
    ],
    "policy": [
        "Cancellation policy",
        "Payment policy",
        "Privacy policy",
    ],
    "procedure": [
        "Clinical procedure",
        "Admin process",
        "Emergency process",
    ],
    "process": [
        "Clinical process",
        "Admin process",
        "Escalation process",
    ],
}

def normalize_question_text(question: str | None) -> str:
    if not question:
        return ""
    return " ".join(question.strip().lower().split())

def is_identity_question(normalized_question: str) -> bool:
    if not normalized_question:
        return False

    identity_phrases = [
        "who are you",
        "what are you",
        "are you ai",
        "are you dental buddy ai",
        "what is dental buddy ai",
    ]

    return any(phrase in normalized_question for phrase in identity_phrases)

def is_help_question(normalized_question: str) -> bool:
    if not normalized_question:
        return False

    help_phrases = [
        "what can you do",
        "how can you help",
        "what can you do for me",
        "how do you help",
        "what do you do",
        "what do you help with",
        "what can i ask",
        "what kinds of questions can i ask",
        "how can you support me",
    ]

    return any(phrase in normalized_question for phrase in help_phrases)

def is_acknowledgement_message(normalized_question: str) -> bool:
    if not normalized_question:
        return False

    acknowledgement_phrases = [
        "yes you have answered it",
        "yes that answered it",
        "that answered it",
        "yes thanks",
        "thanks",
        "thank you",
        "all good",
        "all good thanks",
        "perfect thanks",
        "that helps",
        "that helped",
        "cheers",
        "okay thanks",
        "ok thanks",
        "got it thanks",
    ]

    return any(phrase in normalized_question for phrase in acknowledgement_phrases)

def has_strong_search_anchor(normalized_question: str) -> bool:
    if not normalized_question:
        return False

    if re.search(r"\bstep\s+\d+\b", normalized_question):
        return True

    strong_patterns = [
        r"\bprocedure for\b",
        r"\bprocess for\b",
        r"\bpolicy for\b",
        r"\bprotocol for\b",
        r"\bsteps for\b",
        r"\bwhat is the .* procedure\b",
        r"\bwhat is the .* process\b",
        r"\bwho do i notify after\b",
        r"\bhow do we handle\b",
        r"\bhow do we manage\b",
        r"\bwhat does our .* policy say\b",
    ]

    return any(re.search(pattern, normalized_question) for pattern in strong_patterns)

def needs_clarification(normalized_question: str) -> tuple[bool, str | None]:
    if not normalized_question:
        return True, "missing_topic"

    if has_strong_search_anchor(normalized_question):
        return False, None
    
    words = re.findall(r"\b\w+\b", normalized_question)

    vague_short_questions = {
        "help",
        "issue",
        "problem",
        "emergency",
        "infection",
        "consent",
        "complaint",
        "incident",
        "policy",
        "procedure",
        "process",
        "sterilisation",
    }

    ambiguous_reference_words = {"this", "that", "it", "here", "there", "these", "those"}

    for topic in CLARIFICATION_SUGGESTIONS.keys():
        if (
            normalized_question == topic
            or normalized_question.startswith(f"{topic} ")
            or normalized_question.endswith(f" {topic}")
            or f" {topic} " in normalized_question
            or f"{topic} issue" in normalized_question
            or f"{topic} problem" in normalized_question
            or f"help with {topic}" in normalized_question
            or f"question about {topic}" in normalized_question
        ):
            return True, topic
        
    if len(words) <= 3 and normalized_question in vague_short_questions:
        return True, "missing_topic"

    if len(words) <= 5 and any(word in ambiguous_reference_words for word in words):
        return True, "ambiguous_reference"

    missing_topic_phrases = [
        "what do i do",
        "what happens next",
        "can you help with this",
        "what is the procedure",
        "what is the process",
        "who do i notify",
        "where do i find that",
        "what step is this",
        "what is the next step",
    ]

    if any(phrase in normalized_question for phrase in missing_topic_phrases):
        return True, "missing_topic"

    if "step" in normalized_question and not re.search(r"\bstep\s+\d+\b", normalized_question):
        return True, "missing_step_context"

    return False, None

def build_clarification_response(clarification_key: str | None) -> str:
    if clarification_key == "missing_step_context":
        return """### I need a bit more detail

Which procedure are you asking about?

If you know the step number too, include that as well.

Once you tell me, I can check the clinic documents."""

    if clarification_key == "ambiguous_reference":
        return """### I need a bit more detail

Can you tell me which procedure, policy, or situation you mean?

Once you tell me, I can check the clinic documents."""

    if clarification_key == "missing_topic":
        return """### I need a bit more detail

Can you tell me which procedure, policy, or situation you need help with?

Once you tell me, I can check the clinic documents."""

    suggestions = CLARIFICATION_SUGGESTIONS.get(clarification_key or "")
    if suggestions:
        bullets = "\n".join(f"- {item}" for item in suggestions)
        return f"""### Which one do you mean?

Choose one of these so I can check the right clinic guidance:

{bullets}"""

    return """### I need a bit more detail

Can you tell me which procedure, policy, or situation you need help with?

Once you tell me, I can check the clinic documents."""

def is_strong_document_query(normalized_question: str) -> bool:
    if not normalized_question:
        return False

    if has_strong_search_anchor(normalized_question):
        return True

    strong_domain_terms = [
        "policy",
        "procedure",
        "process",
        "protocol",
        "consent",
        "sterilisation",
        "sterilization",
        "infection",
        "incident",
        "complaint",
        "emergency",
        "form",
        "document",
        "report",
        "notify",
        "steps",
        "needle stick",
        "cross contamination",
        "treatment consent",
        "financial consent",
        "privacy policy",
        "cancellation policy",
        "payment policy",
    ]

    return any(term in normalized_question for term in strong_domain_terms)

def classify_route_with_llm(question: str) -> str:
    try:
        response = azure_client.chat.completions.create(
            model=CHAT_DEPLOYMENT,
            messages=[
                {
                    "role": "system",
                    "content": """
You are classifying a user message for Dental Buddy AI.

Return only one label:
- greeting
- help
- clarification
- document_query

Rules:
- greeting = casual hello, greeting, or light small talk with no real clinic question
- help = the user is asking what Dental Buddy AI can do, how it helps, or what kinds of questions they can ask
- clarification = vague request that lacks enough detail to search clinic documents properly
- document_query = a real clinic, policy, procedure, protocol, process, form, or document-related question

Important:
- If the user includes a real clinic question, return document_query even if the message starts with a greeting
- Be tolerant of typos and casual phrasing
- Return only the label, nothing else
"""
                },
                {
                    "role": "user",
                    "content": question,
                },
            ],
            temperature=0,
        )

        label = (response.choices[0].message.content or "").strip().lower()

        if label in {"greeting", "help", "clarification", "document_query"}:
            return label

    except Exception:
        pass

    return "document_query"

def generate_greeting_response(question: str, clinic_name: str) -> str:
    try:
        response = azure_client.chat.completions.create(
            model=CHAT_DEPLOYMENT,
            messages=[
                {
                    "role": "system",
                    "content": build_greeting_prompt(clinic_name),
                },
                {
                    "role": "user",
                    "content": question,
                },
            ],
            temperature=0.6,
        )

        content = (response.choices[0].message.content or "").strip()
        if content:
            return content

    except Exception:
        pass

    return f"Hey — how can I help at {clinic_name}?"

def generate_identity_response(question: str, clinic_name: str) -> str:
    try:
        response = azure_client.chat.completions.create(
            model=CHAT_DEPLOYMENT,
            messages=[
                {
                    "role": "system",
                    "content": build_identity_prompt(clinic_name),
                },
                {
                    "role": "user",
                    "content": question,
                },
            ],
            temperature=0.5,
        )

        content = (response.choices[0].message.content or "").strip()
        if content:
            return content

    except Exception:
        pass

    return f"I'm Dental Buddy AI for {clinic_name}. I help by looking through your clinic documents and answering questions based on them."

def generate_help_response(question: str, clinic_name: str) -> str:
    try:
        response = azure_client.chat.completions.create(
            model=CHAT_DEPLOYMENT,
            messages=[
                {
                    "role": "system",
                    "content": build_help_prompt(clinic_name),
                },
                {
                    "role": "user",
                    "content": question,
                },
            ],
            temperature=0.5,
        )

        content = (response.choices[0].message.content or "").strip()
        if content:
            return content

    except Exception:
        pass

    return (
        f"I can help answer questions for {clinic_name} using the clinic's documents, "
        f"like procedures, policies, forms, and other clinic guidance."
    )

def build_no_docs_fallback_response() -> str:
    return """### No Relevant Clinic Information Found

I couldn’t find relevant information in the clinic documents to answer that question.

### Next Step

You can try asking in a more specific way, or contact your Practice Manager or Clinic Owner for further guidance."""

def filter_sources_by_question_relevance(question: str, sources: list[dict]) -> list[dict]:
    stop_words = {
        "what", "when", "where", "which", "who", "how", "does", "do", "is", "are",
        "the", "our", "your", "for", "and", "with", "about", "tell"
    }

    question_keywords = [
        word
        for word in re.findall(r"\b\w+\b", question.lower())
        if len(word) > 3 and word not in stop_words
    ]

    if not question_keywords:
        return sources

    filtered_sources = []
    for s in sources:
        searchable_text = (
            f"{s.get('title') or ''} "
            f"{s.get('section') or ''} "
            f"{s.get('text') or ''}"
        ).lower()

        if any(keyword in searchable_text for keyword in question_keywords):
            filtered_sources.append(s)

    return filtered_sources

def build_dynamic_clarification_response(
    *,
    question: str,
    db: Session,
    current_user: CurrentUser,
    top_k: int = 8,
) -> str:
    selected_clinic_id = current_user.selectedClinicId

    effective_level = get_effective_permission_level(
        db,
        user=current_user,
        clinic_id=selected_clinic_id,
    )

    allowed_levels = get_accessible_document_levels(effective_level)
    if not allowed_levels:
        return (
            "### I need a bit more detail\n\n"
            f"Could you tell me a bit more about what you mean by {question.strip()}?"
        )

    qvec = embed_text(question)

    must_filters = [
        qm.FieldCondition(
            key="organisationId",
            match=qm.MatchValue(value=current_user.organisationId),
        ),
        qm.FieldCondition(
            key="status",
            match=qm.MatchValue(value="active"),
        ),
    ]

    role_should_filters = [
        qm.FieldCondition(
            key="roleAccess",
            match=qm.MatchValue(value=level),
        )
        for level in allowed_levels
    ]

    org_membership = get_organisation_membership(
        db,
        user_id=current_user.userId,
        organisation_id=current_user.organisationId,
    )
    org_level = normalize_permission_level(
        org_membership.permissionLevel if org_membership else current_user.organisationPermissionLevel
    )
    is_org_admin = org_level == "admin"

    if is_org_admin:
        if selected_clinic_id:
            scope_should_filters = [
                qm.FieldCondition(
                    key="clinicId",
                    match=qm.MatchValue(value=selected_clinic_id),
                ),
                qm.FieldCondition(
                    key="isShared",
                    match=qm.MatchValue(value=True),
                ),
            ]
        else:
            scope_should_filters = [
                qm.FieldCondition(
                    key="isShared",
                    match=qm.MatchValue(value=True),
                )
            ]
    else:
        selected_clinic_id = require_selected_clinic(current_user)

        clinic_membership = get_clinic_membership(
            db,
            user_id=current_user.userId,
            organisation_id=current_user.organisationId,
            clinic_id=selected_clinic_id,
        )
        if not clinic_membership:
            return (
                "### I need a bit more detail\n\n"
                f"Could you tell me a bit more about what you mean by {question.strip()}?"
            )

        scope_should_filters = [
            qm.FieldCondition(
                key="clinicId",
                match=qm.MatchValue(value=selected_clinic_id),
            ),
            qm.FieldCondition(
                key="isShared",
                match=qm.MatchValue(value=True),
            ),
        ]

    query_filter = qm.Filter(
        must=must_filters,
        should=role_should_filters + scope_should_filters,
    )

    response = qdrant.query_points(
        collection_name=COLLECTION,
        query=qvec,
        limit=top_k,
        query_filter=query_filter,
    )

    valid_documents = {
        d.documentId: d
        for d in db.query(Document)
        .filter(
            Document.organisationId == current_user.organisationId,
            Document.status == "active",
        )
        .all()
    }

    suggestions = []
    seen = set()

    for point in response.points:
        payload = point.payload or {}
        document_id = payload.get("documentId")
        db_doc = valid_documents.get(document_id)

        if not db_doc:
            continue

        if db_doc.status != "active":
            continue

        if db_doc.clinicId != payload.get("clinicId"):
            continue

        try:
            require_document_access(
                db,
                user=current_user,
                document=db_doc,
                required_level="read",
            )
        except HTTPException:
            continue

    title = normalize_line(payload.get("title") or "")
    section = normalize_line(payload.get("section") or "")
    text = normalize_line(payload.get("text") or "")

    candidate_options = []

    if title:
        candidate_options.append(title)

    if section and section.lower() != title.lower():
        candidate_options.append(section)

    for candidate in candidate_options:
        candidate = candidate.strip(" -:\n\t|.")

        if not candidate:
            continue

        candidate_lower = candidate.lower()

        blocked_exact = {
            "metadata",
            "content",
            "procedure",
            "procedures",
            "overview",
            "goal",
            "purpose",
            "scope",
            "reporting",
            "responsibilities",
            "responsibility",
            "definitions",
            "references",
            "follow up",
            "follow-up",
            "documentation",
            "notification",
            "notify",
            "owner",
            "approver",
        }

        blocked_contains = [
            "owner |",
            "approver |",
            "goal |",
            "implementation date",
            "last review date",
            "| practice manager",
            "| business manager",
        ]

        if candidate_lower in blocked_exact:
            continue

        if any(item in candidate_lower for item in blocked_contains):
            continue

        if len(candidate) < 4:
            continue

        if len(candidate.split()) > 8:
            continue

        if candidate_lower in seen:
            continue

        seen.add(candidate_lower)
        suggestions.append(candidate)

        if len(suggestions) >= 4:
            break

    cleaned_question = question.strip().rstrip("?.!")

    if suggestions:
        bullets = "\n".join(f"- {item}" for item in suggestions)
        return (
            "### Did you mean one of these?\n\n"
            f"I found a few related topics for **{cleaned_question}** in the clinic documents:\n\n"
            f"{bullets}\n\n"
            "Pick one, or type your question in a bit more detail."
        )

    return (
        "### I need a bit more detail\n\n"
        f"Could you tell me a bit more about what you mean by **{cleaned_question}**?\n\n"
        "Once I have a bit more detail, I can check the clinic documents."
    )

def build_context_from_sources(sources: list[dict]) -> str:
    return "\n\n".join(
        [
            (
                f"Document: {s['filename']}\n"
                f"Title: {s.get('title') or ''}\n"
                f"Section: {s.get('section') or ''}\n"
                f"Step Number: {s.get('stepNumber') or ''}\n"
                f"{s['text']}"
            )
            for s in sources
        ]
    )

def generate_rag_answer_with_retry(
    reranked_sources: list[dict],
    clinic_name: str,
    question: str,
) -> tuple[str, list[dict], bool]:
    chunk_sets = [
        reranked_sources[:4],
        reranked_sources[:2],
        reranked_sources[:1],
    ]

    for attempt_number, sources_subset in enumerate(chunk_sets, start=1):
        if not sources_subset:
            continue

        try:
            context = build_context_from_sources(sources_subset)

            response = azure_client.chat.completions.create(
                model=CHAT_DEPLOYMENT,
                messages=[
                    {
                        "role": "system",
                        "content": build_rag_prompt(clinic_name),
                    },
                    {
                        "role": "user",
                        "content": f"""
Use the following clinic documents to answer the question.

Documents:
{context}

Question:
{question}
"""
                    }
                ],
                temperature=0.2,
            )

            answer = (response.choices[0].message.content or "").strip()
            if answer:
                return answer, sources_subset, False

        except Exception as e:
            if "content_filter" in str(e).lower():
                print("Azure content filter triggered")
                print({
                    "question": question,
                    "attempt_chunk_count": len(sources_subset),
                    "filenames": [s.get("filename") for s in sources_subset],
                    "document_ids": [s.get("documentId") for s in sources_subset],
                })
                print("Blocked chunk texts:")
                for s in sources_subset:
                    print({
                        "filename": s.get("filename"),
                        "title": s.get("title"),
                        "section": s.get("section"),
                        "text": s.get("text"),
                    })
                continue
            raise

    fallback_answer = """### I couldn’t return that safely just yet

I found related clinic information, but couldn’t generate a final answer from it.

Please try asking in a more specific way, or contact your Practice Manager or Clinic Owner if needed."""

    return fallback_answer, [], True

def generate_acknowledgement_response() -> str:
    responses = [
        "Glad that helped. Let me know if you need anything else.",
        "Great, happy to help. Just let me know if anything else comes up.",
        "No worries, glad that answered it. I'm here if you need anything else.",
    ]
    return random.choice(responses)

def assess_question_route(question: str | None) -> dict:
    normalized_question = normalize_question_text(question)

    if is_identity_question(normalized_question):
        return {
            "route": "identity",
            "normalizedQuestion": normalized_question,
        }

    if is_help_question(normalized_question):
        return {
            "route": "help",
            "normalizedQuestion": normalized_question,
        }

    if is_strong_document_query(normalized_question):
        clarification_needed, clarification_key = needs_clarification(normalized_question)
        if clarification_needed:
            return {
                "route": "clarification",
                "normalizedQuestion": normalized_question,
                "clarificationKey": clarification_key,
            }

        return {
            "route": "document_query",
            "normalizedQuestion": normalized_question,
        }

    llm_route = classify_route_with_llm(question or "")

    if llm_route == "greeting":
        return {
            "route": "greeting",
            "normalizedQuestion": normalized_question,
        }

    if llm_route == "help":
        return {
            "route": "help",
            "normalizedQuestion": normalized_question,
        }

    if is_acknowledgement_message(normalized_question):
        return {
            "route": "acknowledgement",
            "normalizedQuestion": normalized_question,
        }

    if llm_route == "clarification":
        clarification_needed, clarification_key = needs_clarification(normalized_question)
        return {
            "route": "clarification",
            "normalizedQuestion": normalized_question,
            "clarificationKey": clarification_key,
        }

    clarification_needed, clarification_key = needs_clarification(normalized_question)
    if clarification_needed:
        return {
            "route": "clarification",
            "normalizedQuestion": normalized_question,
            "clarificationKey": clarification_key,
        }

    return {
        "route": "document_query",
        "normalizedQuestion": normalized_question,
    }

def append_confirmation(answer: str) -> str:
    if not answer:
        return answer

    return answer + """

### Need anything else?

Did this answer your question, or is there something more specific you’d like help with?
"""

# -------- Models --------
class IngestRequest(BaseModel):
    documentId: str

class ArchiveDocumentRequest(BaseModel):
    documentId: str

class DeleteDocumentRequest(BaseModel):
    documentId: str

class RepairDocumentRequest(BaseModel):
    documentId: str

class BulkRepairDocumentsRequest(BaseModel):
    onlyIssues: bool = True
    limit: int = 50

class ReindexDocumentsRequest(BaseModel):
    onlyNonIndexed: bool = False
    limit: int = 50

class DeleteOrphanedVectorsRequest(BaseModel):
    documentIds: list[str] = Field(default_factory=list)

class OrphanedVectorItem(BaseModel):
    documentId: str
    vectorCount: int

class AskRequest(BaseModel):
    question: str
    topK: int = 10
    conversationId: str | None = None

class ClinicMembershipResponse(BaseModel):
    clinicId: str
    permissionLevel: str

class CurrentUserResponse(BaseModel):
    userId: str
    displayName: str
    email: str | None = None
    username: str | None = None
    accountType: str
    status: str
    organisationId: str
    organisationPermissionLevel: str | None = None
    selectedClinicId: str | None = None
    effectivePermissionLevel: str | None = None
    clinicMemberships: list[ClinicMembershipResponse]

class ClinicAccessInput(BaseModel):
    clinicId: str
    permissionLevel: str  # admin | manage | write | read

class ClinicAccessAssignmentRequest(BaseModel):
    clinicId: str
    permissionLevel: str  # admin | manage | write | read

class CreateUserRequest(BaseModel):
    displayName: str
    accountType: str = "work"  # work | workstation
    email: str | None = None
    username: str | None = None
    password: str | None = None
    scope: str  # organisation | clinic
    permissionLevel: str | None = None  # used for organisation scope
    clinicAccess: list[ClinicAccessAssignmentRequest] = []  # used for clinic scope

class UserListItemResponse(BaseModel):
    userId: str
    displayName: str
    email: str | None = None
    username: str | None = None
    accountType: str
    status: str
    organisationId: str
    organisationPermissionLevel: str | None = None
    clinicMemberships: list[ClinicMembershipResponse]

class BootstrapAdminRequest(BaseModel):
    organisationId: str
    displayName: str
    email: str

class UpdateUserStatusRequest(BaseModel):
    userId: str
    status: str  # active | disabled

class DeleteUserRequest(BaseModel):
    userId: str

class UpdateUserAccessRequest(BaseModel):
    userId: str
    scope: str  # organisation | clinic
    permissionLevel: str | None = None  # used for organisation scope
    clinicAccess: list[ClinicAccessAssignmentRequest] = []  # used for clinic scope

class ClinicListItemResponse(BaseModel):
    clinicId: str
    name: str
    organisationId: str

class ClinicsListResponse(BaseModel):
    clinics: list[ClinicListItemResponse]

class SetPasswordRequest(BaseModel):
    email: str
    password: str

class DocumentConsistencyItem(BaseModel):
    documentId: str
    filename: str
    clinicId: str | None = None
    isShared: bool
    status: str | None = None
    indexStatus: str | None = None
    vectorCount: int
    isConsistent: bool
    issues: list[str]

class DocumentHealthSummaryResponse(BaseModel):
    totalDocuments: int
    activeDocuments: int
    archivedDocuments: int
    indexedDocuments: int
    pendingDocuments: int
    failedDocuments: int
    staleDocuments: int
    inconsistentDocuments: int
    orphanedVectorGroups: int

class AllowedIPEntryInput(BaseModel):
    value: str
    label: str | None = None

class SaveNetworkSettingsRequest(BaseModel):
    mode: str
    entries: list[AllowedIPEntryInput] = []

# -------- Routes --------
@app.get("/health")
def health():
    return {"status": "ok", "service": "dba-mvp-api"}

@app.get("/debug/bootstrap-orgs")
def debug_bootstrap_orgs(db: Session = Depends(get_db)):
    import os
    from database.db import engine

    orgs = db.query(Organisation).all()

    return {
        "cwd": os.getcwd(),
        "engine_url": str(engine.url),
        "organisations": [
            {
                "organisationId": org.organisationId,
                "name": org.name,
            }
            for org in orgs
        ],
    }

@app.get("/users/me", response_model=CurrentUserResponse)
def get_users_me(
    db: Session = Depends(get_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    clinic_memberships = get_clinic_memberships_for_user(
        db,
        user_id=current_user.userId,
        organisation_id=current_user.organisationId,
    )

    effective_permission_level = get_effective_permission_level(
        db,
        user=current_user,
        clinic_id=current_user.selectedClinicId,
    )

    user_record = get_user_by_id(db, current_user.userId)
    if not user_record:
        raise HTTPException(status_code=404, detail="User not found")

    return CurrentUserResponse(
        userId=user_record.userId,
        displayName=user_record.displayName,
        email=user_record.email,
        username=getattr(user_record, "username", None),
        accountType=(getattr(user_record, "accountType", None) or "work"),
        status=user_record.status,
        organisationId=current_user.organisationId,
        organisationPermissionLevel=current_user.organisationPermissionLevel,
        selectedClinicId=current_user.selectedClinicId,
        effectivePermissionLevel=effective_permission_level,
        clinicMemberships=[
            ClinicMembershipResponse(
                clinicId=membership.clinicId,
                permissionLevel=membership.permissionLevel,
            )
            for membership in clinic_memberships
        ],
    )

@app.post("/users/create", response_model=UserListItemResponse)
def create_user(
    req: CreateUserRequest,
    db: Session = Depends(get_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    normalized_display_name = (req.displayName or "").strip()
    normalized_account_type = (req.accountType or "work").strip().lower()
    normalized_email = (req.email or "").strip().lower() or None
    normalized_username = (req.username or "").strip().lower() or None
    raw_password = (req.password or "").strip()

    if not normalized_display_name:
        raise HTTPException(status_code=400, detail="Display name is required")

    if normalized_account_type not in {"work", "workstation"}:
        raise HTTPException(status_code=400, detail="Invalid account type")

    normalized_scope, normalized_permission_level, clinic_access_map = require_user_management_access(
        db,
        actor=current_user,
        scope=req.scope,
        permission_level=req.permissionLevel,
        clinic_access=req.clinicAccess,
    )

    if normalized_account_type == "work":
        if not normalized_email:
            raise HTTPException(status_code=400, detail="Email is required for work accounts")

        existing_user = get_user_by_email(db, normalized_email)
        if existing_user:
            raise HTTPException(status_code=409, detail="A user with this email already exists")

        new_user = User(
            userId=str(uuid.uuid4()),
            email=normalized_email,
            username=None,
            displayName=normalized_display_name,
            accountType="work",
            status="active",
            passwordHash=None,
            mustSetPassword=True,
        )

    else:
        if normalized_scope != "clinic":
            raise HTTPException(
                status_code=400,
                detail="Workstation accounts must use clinic scope",
            )

        if not normalized_username:
            raise HTTPException(status_code=400, detail="Username is required for workstation accounts")

        if not raw_password:
            raise HTTPException(status_code=400, detail="Password is required for workstation accounts")

        validate_password_strength(raw_password)

        if len(raw_password.encode("utf-8")) > 72:
            raise HTTPException(
                status_code=400,
                detail="Password must be 72 bytes or fewer for this MVP auth setup",
            )

        existing_user = get_user_by_username(db, normalized_username)
        if existing_user:
            raise HTTPException(status_code=409, detail="A user with this username already exists")

        if normalized_email:
            existing_email_user = get_user_by_email(db, normalized_email)
            if existing_email_user:
                raise HTTPException(status_code=409, detail="A user with this email already exists")

        if not clinic_access_map:
            raise HTTPException(
                status_code=400,
                detail="Workstation accounts must be assigned to at least one clinic",
            )

        invalid_permissions = [
            level for level in clinic_access_map.values()
            if normalize_permission_level(level) != "read"
        ]
        if invalid_permissions:
            raise HTTPException(
                status_code=400,
                detail="Workstation accounts must be read-only",
            )

        new_user = User(
            userId=str(uuid.uuid4()),
            email=normalized_email,
            username=normalized_username,
            displayName=normalized_display_name,
            accountType="workstation",
            status="active",
            passwordHash=hash_password(raw_password),
            mustSetPassword=False,
        )

    db.add(new_user)
    db.flush()

    if normalized_scope == "organisation":
        org_membership = OrganisationMembership(
            membershipId=str(uuid.uuid4()),
            userId=new_user.userId,
            organisationId=current_user.organisationId,
            permissionLevel=normalized_permission_level,
            createdBy=current_user.userId,
        )
        db.add(org_membership)
    else:
        clinic_memberships = [
            ClinicMembership(
                membershipId=str(uuid.uuid4()),
                userId=new_user.userId,
                organisationId=current_user.organisationId,
                clinicId=clinic_id,
                permissionLevel=clinic_permission_level,
                createdBy=current_user.userId,
            )
            for clinic_id, clinic_permission_level in clinic_access_map.items()
        ]
        db.add_all(clinic_memberships)

    db.commit()

    created_org_membership = get_organisation_membership(
        db,
        user_id=new_user.userId,
        organisation_id=current_user.organisationId,
    )
    created_clinic_memberships = get_clinic_memberships_for_user(
        db,
        user_id=new_user.userId,
        organisation_id=current_user.organisationId,
    )

    return UserListItemResponse(
        userId=new_user.userId,
        displayName=new_user.displayName,
        email=new_user.email,
        username=getattr(new_user, "username", None),
        accountType=(getattr(new_user, "accountType", None) or "work"),
        status=new_user.status,
        organisationId=current_user.organisationId,
        organisationPermissionLevel=(
            created_org_membership.permissionLevel if created_org_membership else None
        ),
        clinicMemberships=[
            ClinicMembershipResponse(
                clinicId=membership.clinicId,
                permissionLevel=membership.permissionLevel,
            )
            for membership in created_clinic_memberships
        ],
    )

@app.post("/users/bootstrap-admin", response_model=UserListItemResponse)
def bootstrap_admin(
    req: BootstrapAdminRequest,
    db: Session = Depends(get_db),
):
    organisation_id = req.organisationId.strip()
    display_name = req.displayName.strip()
    email = req.email.strip().lower()

    if not organisation_id:
        raise HTTPException(status_code=400, detail="Organisation ID is required")

    if not display_name:
        raise HTTPException(status_code=400, detail="Display name is required")

    if not email:
        raise HTTPException(status_code=400, detail="Email is required")

    organisation = (
        db.query(Organisation)
        .filter(Organisation.organisationId == organisation_id)
        .first()
    )
    if not organisation:
        raise HTTPException(status_code=404, detail="Organisation not found")

    if organisation_has_admin(db, organisation_id):
        raise HTTPException(
            status_code=409,
            detail="This organisation already has an admin",
        )

    existing_user = get_user_by_email(db, email)
    if existing_user:
        raise HTTPException(
            status_code=409,
            detail="A user with this email already exists",
        )

    new_user = User(
        userId=str(uuid.uuid4()),
        email=email,
        username=None,
        displayName=display_name,
        status="active",
    )
    db.add(new_user)
    db.flush()

    org_membership = OrganisationMembership(
        membershipId=str(uuid.uuid4()),
        userId=new_user.userId,
        organisationId=organisation_id,
        permissionLevel="admin",
        createdBy="system-bootstrap",
    )
    db.add(org_membership)

    db.commit()

    return UserListItemResponse(
        userId=new_user.userId,
        displayName=new_user.displayName,
        email=new_user.email,
        username=getattr(new_user, "username", None),
        accountType=(getattr(new_user, "accountType", None) or "work"),
        status=new_user.status,
        organisationId=organisation_id,
        organisationPermissionLevel="admin",
        clinicMemberships=[],
    )

@app.get("/users/list", response_model=list[UserListItemResponse])
def list_users(
    db: Session = Depends(get_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    users = get_visible_users_for_actor(
        db,
        actor=current_user,
    )

    results = []
    for user in users:
        org_membership = get_organisation_membership(
            db,
            user_id=user.userId,
            organisation_id=current_user.organisationId,
        )
        clinic_memberships = get_clinic_memberships_for_user(
            db,
            user_id=user.userId,
            organisation_id=current_user.organisationId,
        )

        results.append(
            UserListItemResponse(
                userId=user.userId,
                displayName=user.displayName,
                email=user.email,
                username=getattr(user, "username", None),
                accountType=(getattr(user, "accountType", None) or "work"),
                status=user.status,
                organisationId=current_user.organisationId,
                organisationPermissionLevel=(
                    org_membership.permissionLevel if org_membership else None
                ),
                clinicMemberships=[
                    ClinicMembershipResponse(
                        clinicId=membership.clinicId,
                        permissionLevel=membership.permissionLevel,
                    )
                    for membership in clinic_memberships
                ],
            )
        )

    return results

@app.post("/users/update-status", response_model=UserListItemResponse)
def update_user_status(
    req: UpdateUserStatusRequest,
    db: Session = Depends(get_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    normalized_status = (req.status or "").strip().lower()
    if normalized_status not in {"active", "disabled"}:
        raise HTTPException(status_code=400, detail="Invalid status")

    target_user = get_user_by_id(db, req.userId)
    if not target_user:
        raise HTTPException(status_code=404, detail="User not found")

    target_org_membership = get_organisation_membership(
        db,
        user_id=target_user.userId,
        organisation_id=current_user.organisationId,
    )
    target_clinic_memberships = get_clinic_memberships_for_user(
        db,
        user_id=target_user.userId,
        organisation_id=current_user.organisationId,
    )

    if not target_org_membership and not target_clinic_memberships:
        raise HTTPException(status_code=404, detail="User is not part of this organisation")

    require_user_status_management_access(
        db,
        actor=current_user,
        target_user=target_user,
        next_status=normalized_status,
    )

    target_user.status = normalized_status
    db.commit()

    refreshed_org_membership = get_organisation_membership(
        db,
        user_id=target_user.userId,
        organisation_id=current_user.organisationId,
    )
    refreshed_clinic_memberships = get_clinic_memberships_for_user(
        db,
        user_id=target_user.userId,
        organisation_id=current_user.organisationId,
    )

    return UserListItemResponse(
        userId=target_user.userId,
        displayName=target_user.displayName,
        email=target_user.email,
        username=getattr(target_user, "username", None),
        accountType=(getattr(target_user, "accountType", None) or "work"),
        status=target_user.status,
        organisationId=current_user.organisationId,
        organisationPermissionLevel=(
            refreshed_org_membership.permissionLevel if refreshed_org_membership else None
        ),
        clinicMemberships=[
            ClinicMembershipResponse(
                clinicId=membership.clinicId,
                permissionLevel=membership.permissionLevel,
            )
            for membership in refreshed_clinic_memberships
        ],
    )

@app.post("/users/update-access", response_model=UserListItemResponse)
def update_user_access(
    req: UpdateUserAccessRequest,
    db: Session = Depends(get_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    target_user = get_user_by_id(db, req.userId)
    if not target_user:
        raise HTTPException(status_code=404, detail="User not found")

    target_org_membership = get_organisation_membership(
        db,
        user_id=target_user.userId,
        organisation_id=current_user.organisationId,
    )
    target_clinic_memberships = get_clinic_memberships_for_user(
        db,
        user_id=target_user.userId,
        organisation_id=current_user.organisationId,
    )

    if not target_org_membership and not target_clinic_memberships:
        raise HTTPException(status_code=404, detail="User is not part of this organisation")

    normalized_scope, normalized_permission_level, clinic_access_map = require_user_access_update_management(
        db,
        actor=current_user,
        target_user=target_user,
        scope=req.scope,
        permission_level=req.permissionLevel,
        clinic_access=req.clinicAccess,
    )

    if normalized_scope == "organisation":
        existing_org_membership = get_organisation_membership(
            db,
            user_id=target_user.userId,
            organisation_id=current_user.organisationId,
        )

        if existing_org_membership:
            existing_org_membership.permissionLevel = normalized_permission_level
        else:
            db.add(
                OrganisationMembership
                (
                    membershipId=str(uuid.uuid4()),
                    userId=target_user.userId,
                    organisationId=current_user.organisationId,
                    permissionLevel=normalized_permission_level,
                    createdBy=current_user.userId,
                )
            )

        for membership in target_clinic_memberships:
            db.delete(membership)

    else:
        existing_org_membership = get_organisation_membership(
            db,
            user_id=target_user.userId,
            organisation_id=current_user.organisationId,
        )
        if existing_org_membership:
            db.delete(existing_org_membership)

        existing_by_clinic = {
            membership.clinicId: membership for membership in target_clinic_memberships
        }
        requested_by_clinic = clinic_access_map

        for membership in target_clinic_memberships:
            if membership.clinicId not in requested_by_clinic:
                db.delete(membership)

        for clinic_id, permission_level in requested_by_clinic.items():
            if clinic_id in existing_by_clinic:
                existing_by_clinic[clinic_id].permissionLevel = permission_level
            else:
                db.add(
                    ClinicMembership(
                        membershipId=str(uuid.uuid4()),
                        userId=target_user.userId,
                        organisationId=current_user.organisationId,
                        clinicId=clinic_id,
                        permissionLevel=permission_level,
                        createdBy=current_user.userId,
                    )
                )

    db.commit()

    refreshed_org_membership = get_organisation_membership(
        db,
        user_id=target_user.userId,
        organisation_id=current_user.organisationId,
    )
    refreshed_clinic_memberships = get_clinic_memberships_for_user(
        db,
        user_id=target_user.userId,
        organisation_id=current_user.organisationId,
    )

    return UserListItemResponse(
        userId=target_user.userId,
        displayName=target_user.displayName,
        email=target_user.email,
        username=getattr(target_user, "username", None),
        accountType=(getattr(target_user, "accountType", None) or "work"),
        status=target_user.status,
        organisationId=current_user.organisationId,
        organisationPermissionLevel=(
            refreshed_org_membership.permissionLevel if refreshed_org_membership else None
        ),
        clinicMemberships=[
            ClinicMembershipResponse(
                clinicId=membership.clinicId,
                permissionLevel=membership.permissionLevel,
            )
            for membership in refreshed_clinic_memberships
        ],
    )

@app.post("/users/delete")
def delete_user(
    req: DeleteUserRequest,
    db: Session = Depends(get_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    target_user = get_user_by_id(db, req.userId)
    if not target_user:
        raise HTTPException(status_code=404, detail="User not found")

    target_org_membership = get_organisation_membership(
        db,
        user_id=target_user.userId,
        organisation_id=current_user.organisationId,
    )
    target_clinic_memberships = get_clinic_memberships_for_user(
        db,
        user_id=target_user.userId,
        organisation_id=current_user.organisationId,
    )

    if not target_org_membership and not target_clinic_memberships:
        raise HTTPException(status_code=404, detail="User is not part of this organisation")

    require_user_delete_management_access(
        db,
        actor=current_user,
        target_user=target_user,
    )

    if target_org_membership:
        db.delete(target_org_membership)

    for membership in target_clinic_memberships:
        db.delete(membership)

    deleted_user_id = target_user.userId
    deleted_display_name = target_user.displayName
    deleted_email = target_user.email

    db.delete(target_user)
    db.commit()

    return {
        "status": "deleted",
        "userId": deleted_user_id,
        "displayName": deleted_display_name,
        "email": deleted_email,
    }

@app.get("/clinics/list", response_model=ClinicsListResponse)
def list_clinics(
    db: Session = Depends(get_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    org_membership = get_organisation_membership(
        db,
        user_id=current_user.userId,
        organisation_id=current_user.organisationId,
    )
    org_level = normalize_permission_level(
        org_membership.permissionLevel if org_membership else current_user.organisationPermissionLevel
    )
    is_org_admin = org_level == "admin"

    if is_org_admin:
        clinics = (
            db.query(Clinic)
            .filter(Clinic.organisationId == current_user.organisationId)
            .order_by(Clinic.name.asc())
            .all()
        )
    else:
        clinic_memberships = get_clinic_memberships_for_user(
            db,
            user_id=current_user.userId,
            organisation_id=current_user.organisationId,
        )

        clinic_ids = [membership.clinicId for membership in clinic_memberships]

        if not clinic_ids:
            return ClinicsListResponse(clinics=[])

        clinics = (
            db.query(Clinic)
            .filter(
                Clinic.organisationId == current_user.organisationId,
                Clinic.clinicId.in_(clinic_ids),
            )
            .order_by(Clinic.name.asc())
            .all()
        )

    return ClinicsListResponse(
        clinics=[
            ClinicListItemResponse(
                clinicId=clinic.clinicId,
                name=clinic.name,
                organisationId=clinic.organisationId,
            )
            for clinic in clinics
        ]
    )

@app.post("/docs/upload")
async def upload_doc(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    clinicId: str = Form(""),
    documentType: str = Form("sop"),
    roleAccess: str = Form("read"),
    sourceType: str = Form("internal"),
    sourceUrl: str = Form(""),
    isShared: bool = Form(False),
    db: Session = Depends(get_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    normalized_role_access = normalize_permission_level(roleAccess)
    if not normalized_role_access:
        raise HTTPException(status_code=400, detail="Invalid document access level")

    target_clinic_id = clinicId.strip() if clinicId else None

    if isShared:
        effective_level = get_effective_permission_level(
            db,
            user=current_user,
            clinic_id=target_clinic_id,
        )
        if not effective_level or not can_upload_document(effective_level):
            raise HTTPException(status_code=403, detail="You do not have permission to upload documents")

        if not can_manage_shared_documents(effective_level):
            raise HTTPException(status_code=403, detail="You do not have permission to upload shared documents")

        if not can_grant_document_level(effective_level, normalized_role_access):
            raise HTTPException(status_code=403, detail="You cannot assign this document access level")

        target_clinic_id = None
    else:
        if not target_clinic_id:
            target_clinic_id = require_selected_clinic(current_user)

        effective_level = require_effective_permission(
            db,
            user=current_user,
            clinic_id=target_clinic_id,
            required_level="write",
        )

        if not can_upload_document(effective_level):
            raise HTTPException(status_code=403, detail="You do not have permission to upload documents")

        if not can_grant_document_level(effective_level, normalized_role_access):
            raise HTTPException(status_code=403, detail="You cannot assign this document access level")

    organisation_id = current_user.organisationId
    document_id = str(uuid.uuid4())

    contents = await file.read()

    storage_provider, storage_path, save_path = save_file_to_storage(
        organisation_id=organisation_id,
        clinic_id=None if isShared else target_clinic_id,
        filename=file.filename,
        contents=contents,
    )

    doc = Document(
        documentId=document_id,
        organisationId=organisation_id,
        clinicId=None if isShared else target_clinic_id,
        filename=file.filename,
        documentType=documentType,
        roleAccess=normalized_role_access,
        sourceType=sourceType,
        sourceUrl=sourceUrl if sourceUrl else None,
        isShared=isShared,
        isCurrentVerified=True,
        uploadedBy=current_user.displayName,
        status="active",
        storageProvider=storage_provider,
        storagePath=storage_path,
        indexStatus="pending",
        indexError=None,
        indexedAt=None,
    )

    db.add(doc)

    create_document_audit_log(
        db,
        action="uploaded",
        organisation_id=organisation_id,
        clinic_id=None if isShared else target_clinic_id,
        performed_by=current_user.displayName,
        filename=file.filename,
        document_id=document_id,
        notes="Document uploaded and queued for processing",
    )

    db.commit()

    background_tasks.add_task(process_document_background, document_id)

    return {
        "status": "uploaded",
        "processingStatus": "processing",
        "documentId": document_id,
        "organisationId": organisation_id,
        "clinicId": target_clinic_id,
        "filename": file.filename,
        "path": str(save_path),
        "sizeBytes": len(contents),
    }

@app.get("/docs/download/{document_id}")
def download_doc(
    document_id: str,
    db: Session = Depends(get_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    document = db.query(Document).filter(Document.documentId == document_id).first()

    if not document:
        return {"status": "error", "message": "Document not found"}

    effective_level = require_document_access(
        db,
        user=current_user,
        document=document,
        required_level="read",
    )

    if not can_download_document(effective_level):
        raise HTTPException(status_code=403, detail="You do not have permission to download this document")

    doc_path = get_document_absolute_path(document)

    if not doc_path.exists():
        return {"status": "error", "message": "File not found on disk"}

    return FileResponse(
        path=str(doc_path),
        filename=document.filename,
        media_type="application/octet-stream",
    )

@app.post("/docs/replace")
async def replace_doc(
    background_tasks: BackgroundTasks,
    oldDocumentId: str = Form(...),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    old_doc = db.query(Document).filter(Document.documentId == oldDocumentId).first()

    if not old_doc:
        return {"status": "error", "message": "Original document not found"}

    effective_level = require_document_access(
        db,
        user=current_user,
        document=old_doc,
        required_level="write",
    )

    if not can_replace_document(effective_level):
        raise HTTPException(status_code=403, detail="You do not have permission to replace this document")

    new_document_id = str(uuid.uuid4())

    contents = await file.read()

    storage_provider, storage_path, save_path = save_file_to_storage(
        organisation_id=old_doc.organisationId,
        clinic_id=old_doc.clinicId,
        filename=file.filename,
        contents=contents,
    )

    new_doc = Document(
        documentId=new_document_id,
        organisationId=old_doc.organisationId,
        clinicId=old_doc.clinicId,
        filename=file.filename,
        documentType=old_doc.documentType,
        roleAccess=old_doc.roleAccess,
        sourceType=old_doc.sourceType,
        sourceUrl=old_doc.sourceUrl,
        isShared=old_doc.isShared,
        isCurrentVerified=True,
        uploadedBy=current_user.displayName,
        status="active",
        storageProvider=storage_provider,
        storagePath=storage_path,
        indexStatus="pending",
        indexError=None,
        indexedAt=None,
    )

    db.add(new_doc)

    old_doc.status = "archived"
    old_doc.indexStatus = "stale"

    create_document_audit_log(
        db,
        action="replaced",
        organisation_id=old_doc.organisationId,
        clinic_id=old_doc.clinicId,
        performed_by=current_user.displayName,
        filename=file.filename,
        document_id=new_document_id,
        old_document_id=old_doc.documentId,
        new_document_id=new_document_id,
        notes=f"Replaced document {old_doc.filename}",
    )

    delete_document_vectors(old_doc.documentId)

    db.commit()

    background_tasks.add_task(process_document_background, new_document_id)

    return {
        "status": "replaced",
        "oldDocumentId": old_doc.documentId,
        "newDocumentId": new_document_id,
        "filename": file.filename,
        "processingStatus": "processing",
    }

@app.get("/docs/list")
def list_docs(
    db: Session = Depends(get_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    selected_clinic_id = current_user.selectedClinicId
    allowed_levels = get_allowed_document_levels_for_user(
        db,
        user=current_user,
        clinic_id=selected_clinic_id,
    )

    if not allowed_levels:
        return {"documents": []}

    org_membership = get_organisation_membership(
        db,
        user_id=current_user.userId,
        organisation_id=current_user.organisationId,
    )
    org_level = normalize_permission_level(
        org_membership.permissionLevel if org_membership else current_user.organisationPermissionLevel
    )
    is_org_admin = org_level == "admin"

    query = db.query(Document).filter(
        Document.organisationId == current_user.organisationId,
        Document.roleAccess.in_(allowed_levels),
    )

    if is_org_admin:
        if selected_clinic_id:
            query = query.filter(
                (Document.clinicId == selected_clinic_id)
                | (Document.isShared == True)
            )
        else:
            query = query.filter(Document.isShared == True)
    else:
        if not selected_clinic_id:
            return {"documents": []}

        clinic_membership = get_clinic_membership(
            db,
            user_id=current_user.userId,
            organisation_id=current_user.organisationId,
            clinic_id=selected_clinic_id,
        )

        if not clinic_membership:
            return {"documents": []}

        query = query.filter(
            (Document.clinicId == selected_clinic_id)
            | (Document.isShared == True)
        )

    docs = query.order_by(Document.uploadedAt.desc()).all()

    return {
        "documents": [
            {
                "documentId": d.documentId,
                "filename": d.filename,
                "documentType": d.documentType,
                "roleAccess": d.roleAccess,
                "sourceType": d.sourceType,
                "isShared": d.isShared,
                "status": d.status,
                "indexStatus": d.indexStatus,
                "indexedAt": d.indexedAt,
                "uploadedAt": d.uploadedAt,
                "readiness": d.readiness,
                "readinessNotes": d.readinessNotes,
            }
            for d in docs
        ]
    }

@app.post("/docs/orphaned-vectors/delete")
def delete_orphaned_vectors(
    req: DeleteOrphanedVectorsRequest,
    db: Session = Depends(get_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    selected_clinic_id = current_user.selectedClinicId
    effective_level = get_effective_permission_level(
        db,
        user=current_user,
        clinic_id=selected_clinic_id,
    )

    if not effective_level or not can_view_activity(effective_level):
        raise HTTPException(status_code=403, detail="You do not have permission to delete orphaned vectors")

    orphaned = get_orphaned_vectors_for_org(db, current_user.organisationId)
    orphaned_ids = {item["documentId"] for item in orphaned}

    target_ids = set(req.documentIds or [])
    if not target_ids:
        target_ids = orphaned_ids

    deleted = []
    skipped = []

    for document_id in target_ids:
        if document_id not in orphaned_ids:
            skipped.append(document_id)
            continue

        delete_document_vectors(document_id)
        deleted.append(
            {
                "documentId": document_id,
                "vectorCountAfterDelete": count_document_vectors(document_id),
            }
        )

    return {
        "status": "completed",
        "deleted": deleted,
        "skipped": skipped,
    }

@app.get("/docs/vector-count/{document_id}")
def get_document_vector_count(
    document_id: str,
    db: Session = Depends(get_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    document = db.query(Document).filter(Document.documentId == document_id).first()

    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    require_document_access(
        db,
        user=current_user,
        document=document,
        required_level="read",
    )

    return {
        "documentId": document.documentId,
        "filename": document.filename,
        "vectorCount": count_document_vectors(document.documentId),
        "status": document.status,
        "indexStatus": document.indexStatus,
    }

@app.get("/docs/consistency-check", response_model=list[DocumentConsistencyItem])
def check_documents_consistency(
    onlyIssues: bool = False,
    db: Session = Depends(get_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    selected_clinic_id = current_user.selectedClinicId
    effective_level = get_effective_permission_level(
        db,
        user=current_user,
        clinic_id=selected_clinic_id,
    )

    if not effective_level or not can_view_activity(effective_level):
        raise HTTPException(status_code=403, detail="You do not have permission to view document consistency")

    documents = get_documents_in_scope_for_current_user(
        db,
        current_user=current_user,
    )

    results = [evaluate_document_consistency(document) for document in documents]

    if onlyIssues:
        results = [item for item in results if not item["isConsistent"]]

    return results

@app.post("/docs/repair")
def repair_document(
    req: RepairDocumentRequest,
    db: Session = Depends(get_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    document = db.query(Document).filter(Document.documentId == req.documentId).first()

    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    selected_clinic_id = current_user.selectedClinicId
    effective_level = get_effective_permission_level(
        db,
        user=current_user,
        clinic_id=selected_clinic_id,
    )

    if not effective_level or not can_view_activity(effective_level):
        raise HTTPException(status_code=403, detail="You do not have permission to repair document consistency")

    require_document_access(
        db,
        user=current_user,
        document=document,
        required_level="manage",
    )

    result = repair_document_by_id(req.documentId)
    return result

@app.post("/docs/repair/bulk")
def bulk_repair_documents(
    req: BulkRepairDocumentsRequest,
    db: Session = Depends(get_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    selected_clinic_id = current_user.selectedClinicId
    effective_level = get_effective_permission_level(
        db,
        user=current_user,
        clinic_id=selected_clinic_id,
    )

    if not effective_level or not can_view_activity(effective_level):
        raise HTTPException(status_code=403, detail="You do not have permission to bulk repair document consistency")

    documents = get_documents_in_scope_for_current_user(
        db,
        current_user=current_user,
    )

    safe_limit = max(1, min(req.limit, 200))

    checked = 0
    repaired = 0
    skipped = 0
    results = []

    for document in documents:
        if checked >= safe_limit:
            break

        checked += 1
        consistency = evaluate_document_consistency(document)

        if req.onlyIssues and consistency["isConsistent"]:
            skipped += 1
            continue

        try:
            require_document_access(
                db,
                user=current_user,
                document=document,
                required_level="manage",
            )
        except HTTPException:
            skipped += 1
            continue

        repair_result = repair_document_by_id(document.documentId)
        results.append(
            {
                "documentId": document.documentId,
                "filename": document.filename,
                "consistency": consistency,
                "repairResult": repair_result,
            }
        )

        if repair_result.get("status") == "repaired":
            repaired += 1

    return {
        "status": "completed",
        "checked": checked,
        "repaired": repaired,
        "skipped": skipped,
        "limit": safe_limit,
        "onlyIssues": req.onlyIssues,
        "results": results,
    }

@app.get("/docs/orphaned-vectors", response_model=list[OrphanedVectorItem])
def list_orphaned_vectors(
    db: Session = Depends(get_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    selected_clinic_id = current_user.selectedClinicId
    effective_level = get_effective_permission_level(
        db,
        user=current_user,
        clinic_id=selected_clinic_id,
    )

    if not effective_level or not can_view_activity(effective_level):
        raise HTTPException(status_code=403, detail="You do not have permission to view orphaned vectors")

    return get_orphaned_vectors_for_org(db, current_user.organisationId)

@app.get("/docs/health-summary", response_model=DocumentHealthSummaryResponse)
def get_document_health_summary(
    db: Session = Depends(get_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    selected_clinic_id = current_user.selectedClinicId
    effective_level = get_effective_permission_level(
        db,
        user=current_user,
        clinic_id=selected_clinic_id,
    )

    if not effective_level or not can_view_activity(effective_level):
        raise HTTPException(status_code=403, detail="You do not have permission to view document health summary")

    return build_document_health_summary(
        db,
        current_user=current_user,
    )

@app.post("/docs/reindex")
def reindex_documents(
    req: ReindexDocumentsRequest,
    db: Session = Depends(get_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    selected_clinic_id = current_user.selectedClinicId
    effective_level = get_effective_permission_level(
        db,
        user=current_user,
        clinic_id=selected_clinic_id,
    )

    if not effective_level or not can_view_activity(effective_level):
        raise HTTPException(status_code=403, detail="You do not have permission to reindex documents")

    return reindex_documents_in_scope(
        db,
        current_user=current_user,
        only_non_indexed=req.onlyNonIndexed,
        limit=req.limit,
    )

@app.get("/docs/audit/list")
def list_document_audit_logs(
    limit: int = 50,
    db: Session = Depends(get_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    selected_clinic_id = current_user.selectedClinicId
    effective_level = get_effective_permission_level(
        db,
        user=current_user,
        clinic_id=selected_clinic_id,
    )

    if not effective_level or not can_view_activity(effective_level):
        raise HTTPException(status_code=403, detail="You do not have permission to view activity")

    safe_limit = max(1, min(limit, 200))

    query = db.query(DocumentAuditLog).filter(
        DocumentAuditLog.organisationId == current_user.organisationId,
    )

    org_membership = get_organisation_membership(
        db,
        user_id=current_user.userId,
        organisation_id=current_user.organisationId,
    )
    org_level = normalize_permission_level(
        org_membership.permissionLevel if org_membership else current_user.organisationPermissionLevel
    )
    is_org_admin = org_level == "admin"

    if is_org_admin:
        if selected_clinic_id:
            query = query.filter(
                (DocumentAuditLog.clinicId == selected_clinic_id)
                | (DocumentAuditLog.clinicId.is_(None))
            )
    else:
        selected_clinic_id = require_selected_clinic(current_user)

        clinic_membership = get_clinic_membership(
            db,
            user_id=current_user.userId,
            organisation_id=current_user.organisationId,
            clinic_id=selected_clinic_id,
        )
        if not clinic_membership:
            raise HTTPException(status_code=403, detail="You do not have access to this clinic")

        query = query.filter(
            (DocumentAuditLog.clinicId == selected_clinic_id)
            | (DocumentAuditLog.clinicId.is_(None))
        )

    logs = query.order_by(desc(DocumentAuditLog.performedAt)).limit(safe_limit).all()

    return {
        "auditLogs": [
            {
                "auditId": log.auditId,
                "documentId": log.documentId,
                "oldDocumentId": log.oldDocumentId,
                "newDocumentId": log.newDocumentId,
                "organisationId": log.organisationId,
                "clinicId": log.clinicId,
                "action": log.action,
                "performedBy": log.performedBy,
                "performedAt": log.performedAt,
                "filename": log.filename,
                "notes": log.notes,
            }
            for log in logs
        ]
    }

@app.post("/docs/archive")
def archive_doc(
    req: ArchiveDocumentRequest,
    db: Session = Depends(get_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    document = db.query(Document).filter(Document.documentId == req.documentId).first()

    if not document:
        return {"status": "error", "message": "Document not found"}

    effective_level = require_document_access(
        db,
        user=current_user,
        document=document,
        required_level="manage",
    )

    if not can_archive_document(effective_level):
        raise HTTPException(status_code=403, detail="You do not have permission to archive this document")

    try:
        document.status = "archived"

        create_document_audit_log(
            db,
            action="archived",
            organisation_id=document.organisationId,
            clinic_id=document.clinicId,
            performed_by=current_user.displayName,
            filename=document.filename,
            document_id=document.documentId,
            notes="Document archived",
        )

        db.commit()

        set_document_vectors_status(document.documentId, "archived")

        return {
            "status": "archived",
            "documentId": document.documentId,
            "filename": document.filename,
        }

    except Exception as e:
        db.rollback()
        return {
            "status": "error",
            "message": str(e),
        }

@app.post("/docs/restore")
def restore_doc(
    req: ArchiveDocumentRequest,
    db: Session = Depends(get_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    document = db.query(Document).filter(Document.documentId == req.documentId).first()

    if not document:
        return {"status": "error", "message": "Document not found"}

    effective_level = require_document_access(
        db,
        user=current_user,
        document=document,
        required_level="manage",
    )

    if not can_restore_document(effective_level):
        raise HTTPException(status_code=403, detail="You do not have permission to restore this document")

    try:
        document.status = "active"

        create_document_audit_log(
            db,
            action="restored",
            organisation_id=document.organisationId,
            clinic_id=document.clinicId,
            performed_by=current_user.displayName,
            filename=document.filename,
            document_id=document.documentId,
            notes="Document restored",
        )

        db.commit()

        set_document_vectors_status(document.documentId, "active")

        return {
            "status": "restored",
            "documentId": document.documentId,
            "filename": document.filename,
        }

    except Exception as e:
        db.rollback()
        return {
            "status": "error",
            "message": str(e),
        }

@app.post("/docs/delete")
def delete_doc(
    req: DeleteDocumentRequest,
    db: Session = Depends(get_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    document = db.query(Document).filter(Document.documentId == req.documentId).first()

    if not document:
        return {"status": "error", "message": "Document not found"}

    effective_level = require_document_access(
        db,
        user=current_user,
        document=document,
        required_level="manage",
    )

    if not can_delete_document(effective_level):
        raise HTTPException(status_code=403, detail="You do not have permission to delete this document")

    try:
        delete_document_vectors(document.documentId)

        delete_document_from_storage(document)

        filename = document.filename
        document_id = document.documentId

        create_document_audit_log(
            db,
            action="deleted",
            organisation_id=document.organisationId,
            clinic_id=document.clinicId,
            performed_by=current_user.displayName,
            filename=document.filename,
            document_id=document.documentId,
            notes="Document deleted",
        )

        db.delete(document)
        db.commit()

        return {
            "status": "deleted",
            "documentId": document_id,
            "filename": filename,
        }

    except Exception as e:
        db.rollback()
        return {
            "status": "error",
            "message": str(e),
        }

@app.post("/docs/ingest")
def ingest_doc(req: IngestRequest):
    return process_document_by_id(req.documentId)

@app.post("/ask")
def ask(
    req: AskRequest,
    db: Session = Depends(get_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    selected_clinic_id = current_user.selectedClinicId
    effective_level = get_effective_permission_level(
        db,
        user=current_user,
        clinic_id=selected_clinic_id,
    )

    if not effective_level or not has_level(effective_level, "read"):
        raise HTTPException(status_code=403, detail="You do not have permission to ask questions in this clinic")
    
    clinic_name = get_prompt_clinic_name(db, current_user)

    allowed_levels = get_accessible_document_levels(effective_level)
    if not allowed_levels:
        answer = "I don’t have access to any documents for this clinic."

        create_ask_log(
            db,
            organisation_id=current_user.organisationId,
            clinic_id=selected_clinic_id,
            user_id=current_user.userId,
            conversation_id=req.conversationId,
            question=req.question,
            outcome_status="no_relevant_docs",
            failure_reason="no_accessible_document_levels",
        )
        db.commit()

        return {
            "organisationId": current_user.organisationId,
            "clinicId": selected_clinic_id,
            "conversationId": req.conversationId,
            "question": req.question,
            "answer": answer,
            "sources": [],
            "responseType": "fallback_no_docs",
            "outcomeStatus": "no_relevant_docs",
            "failureReason": "no_accessible_document_levels",
        }
    
    route_result = assess_question_route(req.question)

    if route_result["route"] == "identity":
        answer = generate_identity_response(req.question, clinic_name)

        create_ask_log(
            db,
            organisation_id=current_user.organisationId,
            clinic_id=selected_clinic_id,
            user_id=current_user.userId,
            conversation_id=req.conversationId,
            question=req.question,
            outcome_status="answered",
            failure_reason=None,
        )
        db.commit()

        return {
            "organisationId": current_user.organisationId,
            "clinicId": selected_clinic_id,
            "conversationId": req.conversationId,
            "question": req.question,
            "answer": answer,
            "sources": [],
            "responseType": "identity",
            "outcomeStatus": "answered",
            "failureReason": None,
        }

    if route_result["route"] == "help":
        answer = generate_help_response(req.question, clinic_name)

        create_ask_log(
            db,
            organisation_id=current_user.organisationId,
            clinic_id=selected_clinic_id,
            user_id=current_user.userId,
            conversation_id=req.conversationId,
            question=req.question,
            outcome_status="answered",
            failure_reason=None,
        )
        db.commit()

        return {
            "organisationId": current_user.organisationId,
            "clinicId": selected_clinic_id,
            "conversationId": req.conversationId,
            "question": req.question,
            "answer": answer,
            "sources": [],
            "responseType": "help",
            "outcomeStatus": "answered",
            "failureReason": None,
        }
    
    if route_result["route"] == "greeting":
        answer = generate_greeting_response(req.question, clinic_name)

        create_ask_log(
            db,
            organisation_id=current_user.organisationId,
            clinic_id=selected_clinic_id,
            user_id=current_user.userId,
            conversation_id=req.conversationId,
            question=req.question,
            outcome_status="answered",
            failure_reason=None,
        )
        db.commit()

        return {
            "organisationId": current_user.organisationId,
            "clinicId": selected_clinic_id,
            "conversationId": req.conversationId,
            "question": req.question,
            "answer": answer,
            "sources": [],
            "responseType": "greeting",
            "outcomeStatus": "answered",
            "failureReason": None,
        }

    if route_result["route"] == "acknowledgement":
        answer = generate_acknowledgement_response()

        create_ask_log(
            db,
            organisation_id=current_user.organisationId,
            clinic_id=selected_clinic_id,
            user_id=current_user.userId,
            conversation_id=req.conversationId,
            question=req.question,
            outcome_status="answered",
            failure_reason=None,
        )
        db.commit()

        return {
            "organisationId": current_user.organisationId,
            "clinicId": selected_clinic_id,
            "conversationId": req.conversationId,
            "question": req.question,
            "answer": answer,
            "sources": [],
            "responseType": "acknowledgement",
            "outcomeStatus": "answered",
            "failureReason": None,
        }

    if route_result["route"] == "clarification":
        answer = build_dynamic_clarification_response(
            question=req.question,
            db=db,
            current_user=current_user,
        )

        create_ask_log(
            db,
            organisation_id=current_user.organisationId,
            clinic_id=selected_clinic_id,
            user_id=current_user.userId,
            conversation_id=req.conversationId,
            question=req.question,
            outcome_status="answered",
            failure_reason=None,
        )
        db.commit()

        return {
            "organisationId": current_user.organisationId,
            "clinicId": selected_clinic_id,
            "conversationId": req.conversationId,
            "question": req.question,
            "answer": answer,
            "sources": [],
            "responseType": "clarification",
            "outcomeStatus": "answered",
            "failureReason": None,
        }
        answer = route_result["answer"]

        create_ask_log(
            db,
            organisation_id=current_user.organisationId,
            clinic_id=selected_clinic_id,
            user_id=current_user.userId,
            conversation_id=req.conversationId,
            question=req.question,
            outcome_status="answered",
            failure_reason=None,
        )
        db.commit()

        return {
            "organisationId": current_user.organisationId,
            "clinicId": selected_clinic_id,
            "conversationId": req.conversationId,
            "question": req.question,
            "answer": answer,
            "sources": [],
            "responseType": "clarification",
            "outcomeStatus": "answered",
            "failureReason": None,
        }

    qvec = embed_text(req.question)

    must_filters = [
        qm.FieldCondition(
            key="organisationId",
            match=qm.MatchValue(value=current_user.organisationId),
        ),
        qm.FieldCondition(
            key="status",
            match=qm.MatchValue(value="active"),
        ),
    ]

    role_should_filters = [
        qm.FieldCondition(
            key="roleAccess",
            match=qm.MatchValue(value=level),
        )
        for level in allowed_levels
    ]

    org_membership = get_organisation_membership(
        db,
        user_id=current_user.userId,
        organisation_id=current_user.organisationId,
    )
    org_level = normalize_permission_level(
        org_membership.permissionLevel if org_membership else current_user.organisationPermissionLevel
    )
    is_org_admin = org_level == "admin"

    if is_org_admin:
        if selected_clinic_id:
            scope_should_filters = [
                qm.FieldCondition(
                    key="clinicId",
                    match=qm.MatchValue(value=selected_clinic_id),
                ),
                qm.FieldCondition(
                    key="isShared",
                    match=qm.MatchValue(value=True),
                ),
            ]
        else:
            scope_should_filters = [
                qm.FieldCondition(
                    key="isShared",
                    match=qm.MatchValue(value=True),
                )
            ]
    else:
        selected_clinic_id = require_selected_clinic(current_user)

        clinic_membership = get_clinic_membership(
            db,
            user_id=current_user.userId,
            organisation_id=current_user.organisationId,
            clinic_id=selected_clinic_id,
        )
        if not clinic_membership:
            raise HTTPException(status_code=403, detail="You do not have access to this clinic")

        scope_should_filters = [
            qm.FieldCondition(
                key="clinicId",
                match=qm.MatchValue(value=selected_clinic_id),
            ),
            qm.FieldCondition(
                key="isShared",
                match=qm.MatchValue(value=True),
            ),
        ]

    query_filter = qm.Filter(
        must=must_filters,
        should=role_should_filters + scope_should_filters,
    )

    response = qdrant.query_points(
        collection_name=COLLECTION,
        query=qvec,
        limit=req.topK,
        query_filter=query_filter,
    )

    sources = []
    valid_documents = {
        d.documentId: d
        for d in db.query(Document)
        .filter(
            Document.organisationId == current_user.organisationId,
            Document.status == "active",
        )
        .all()
    }

    for point in response.points:
        payload = point.payload or {}
        document_id = payload.get("documentId")

        # 🚨 Step 1 — Validate against DB
        db_doc = valid_documents.get(document_id)

        if not db_doc:
            continue  # ❌ skip ghost doc

        if db_doc.status != "active":
            continue  # ❌ skip archived/deleted

        # 🚨 Step 2 — Validate clinic match
        if db_doc.clinicId != payload.get("clinicId"):
            continue  # ❌ stale/mismatched payload

        # 🚨 Step 3 — Build source safely
        source = {
            "documentId": db_doc.documentId,
            "filename": db_doc.filename,
            "documentType": db_doc.documentType,
            "roleAccess": db_doc.roleAccess,
            "sourceType": db_doc.sourceType,
            "chunkIndex": payload.get("chunkIndex"),
            "title": payload.get("title"),
            "section": payload.get("section"),
            "stepNumber": payload.get("stepNumber"),
            "score": point.score,
            "text": payload.get("text"),
            "clinicId": db_doc.clinicId,
            "isShared": db_doc.isShared,
        }

        try:
            require_document_access(
                db,
                user=current_user,
                document=db_doc,
                required_level="read",
            )
            sources.append(source)
        except HTTPException:
            continue

    sources = sorted(sources, key=lambda x: x["score"], reverse=True)
    sources = apply_question_heuristics(req.question, sources)
    print("----- DEBUG: RAW SOURCES -----")
    for s in sources[:10]:
        print({
            "filename": s.get("filename"),
            "title": s.get("title"),
            "section": s.get("section"),
            "score": s.get("score"),
            "text": s.get("text"),
        })
    print("--------------------------------")
    reranked_sources = rerank_sources(req.question, sources, max_sources=4)

    print("----- DEBUG: RERANKED SOURCES -----")
    for s in reranked_sources:
        print({
            "documentId": s.get("documentId"),
            "filename": s.get("filename"),
            "clinicId": s.get("clinicId"),
            "isShared": s.get("isShared"),
        })
    print("------------------------------------")

    if not reranked_sources:
        answer = build_no_docs_fallback_response()

        create_ask_log(
            db,
            organisation_id=current_user.organisationId,
            clinic_id=selected_clinic_id,
            user_id=current_user.userId,
            conversation_id=req.conversationId,
            question=req.question,
            outcome_status="no_relevant_docs",
            failure_reason="no_accessible_matching_sources",
        )
        db.commit()

        return {
            "organisationId": current_user.organisationId,
            "clinicId": selected_clinic_id,
            "conversationId": req.conversationId,
            "question": req.question,
            "answer": answer,
            "sources": [],
            "responseType": "fallback_no_docs",
            "outcomeStatus": "no_relevant_docs",
            "failureReason": "no_accessible_matching_sources",
        }

    answer, used_sources, was_filtered_fallback = generate_rag_answer_with_retry(
        reranked_sources=reranked_sources,
        clinic_name=clinic_name,
        question=req.question,
    )

    if was_filtered_fallback:
        create_ask_log(
            db,
            organisation_id=current_user.organisationId,
            clinic_id=selected_clinic_id,
            user_id=current_user.userId,
            conversation_id=req.conversationId,
            question=req.question,
            outcome_status="model_error",
            failure_reason="content_filter_retry_exhausted",
        )
        db.commit()

        return {
            "organisationId": current_user.organisationId,
            "clinicId": selected_clinic_id,
            "conversationId": req.conversationId,
            "question": req.question,
            "answer": answer,
            "sources": [],
            "responseType": "document_answer_filtered",
            "outcomeStatus": "model_error",
            "failureReason": "content_filter_retry_exhausted",
        }

    create_ask_log(
        db,
        organisation_id=current_user.organisationId,
        clinic_id=selected_clinic_id,
        user_id=current_user.userId,
        conversation_id=req.conversationId,
        question=req.question,
        outcome_status="answered",
        failure_reason=None,
    )
    db.commit()

    answer = append_confirmation(answer)

    return {
        "organisationId": current_user.organisationId,
        "clinicId": selected_clinic_id,
        "conversationId": req.conversationId,
        "question": req.question,
        "answer": answer,
        "sources": used_sources,
        "responseType": "document_answer",
        "outcomeStatus": "answered",
        "failureReason": None,
    }

@app.post("/auth/set-password")
def set_password(
    req: SetPasswordRequest,
    db: Session = Depends(get_db),
):
    normalized_email = req.email.strip().lower()
    password = (req.password or "").strip()

    if not normalized_email or not password:
        raise HTTPException(status_code=400, detail="Email and password are required")

    if len(password.encode("utf-8")) > 72:
        raise HTTPException(
            status_code=400,
            detail="Password must be 72 bytes or fewer for this MVP auth setup",
        )

    user = get_user_by_email(db, normalized_email)
    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    user.passwordHash = hash_password(password)
    db.commit()

    return {"status": "ok", "message": "Password set"}

@app.post("/auth/login", response_model=LoginResponse)
def login(req: LoginRequest, db: Session = Depends(get_db)):
    user = get_user_by_email(db, req.email)

    if not user or not user.passwordHash:
        raise HTTPException(status_code=401, detail="Invalid email or password")

    account_type = (getattr(user, "accountType", None) or "work").strip().lower()
    if account_type != "work":
        raise HTTPException(status_code=403, detail="Use the correct login method for this account")

    if user.status != "active":
        raise HTTPException(status_code=403, detail="User is not active")

    if not verify_password(req.password, user.passwordHash):
        raise HTTPException(status_code=401, detail="Invalid email or password")

    org_membership = (
        db.query(OrganisationMembership)
        .filter(OrganisationMembership.userId == user.userId)
        .first()
    )

    clinic_memberships = (
        db.query(ClinicMembership)
        .filter(ClinicMembership.userId == user.userId)
        .all()
    )

    organisation_id = None
    organisation_permission_level = None

    if org_membership:
        organisation_id = org_membership.organisationId
        organisation_permission_level = normalize_permission_level(org_membership.permissionLevel)
    elif clinic_memberships:
        organisation_id = clinic_memberships[0].organisationId
    else:
        raise HTTPException(
            status_code=403,
            detail="User is not assigned to an organisation or clinic",
        )

    user.lastLoginAt = datetime.utcnow()
    db.commit()

    access_token = create_access_token(
        user_id=user.userId,
        organisation_id=organisation_id,
        account_type=account_type,
    )

    return LoginResponse(
        accessToken=access_token,
        mustSetPassword=bool(user.mustSetPassword),
        user=AuthUserPayload(
            userId=user.userId,
            displayName=user.displayName,
            email=user.email,
            username=getattr(user, "username", None),
            accountType=account_type,
            organisationId=organisation_id,
            organisationPermissionLevel=organisation_permission_level,
            clinicMemberships=[
                ClinicMembershipPayload(
                    clinicId=membership.clinicId,
                    permissionLevel=normalize_permission_level(membership.permissionLevel),
                )
                for membership in clinic_memberships
            ],
        ),
    )

@app.post("/auth/login-workstation", response_model=LoginResponse)
def login_workstation(req: WorkstationLoginRequest, db: Session = Depends(get_db)):
    user = get_user_by_username(db, req.username)

    if not user or not user.passwordHash:
        raise HTTPException(status_code=401, detail="Invalid username or password")

    account_type = (getattr(user, "accountType", None) or "work").strip().lower()
    if account_type != "workstation":
        raise HTTPException(status_code=403, detail="Use the correct login method for this account")

    if user.status != "active":
        raise HTTPException(status_code=403, detail="User is not active")

    if not verify_password(req.password, user.passwordHash):
        raise HTTPException(status_code=401, detail="Invalid username or password")

    org_membership = (
        db.query(OrganisationMembership)
        .filter(OrganisationMembership.userId == user.userId)
        .first()
    )

    clinic_memberships = (
        db.query(ClinicMembership)
        .filter(ClinicMembership.userId == user.userId)
        .all()
    )

    if org_membership:
        raise HTTPException(
            status_code=403,
            detail="Workstation accounts must not use organisation-level access",
        )

    if not clinic_memberships:
        raise HTTPException(
            status_code=403,
            detail="Workstation account is not assigned to a clinic",
        )

    invalid_permissions = [
        membership.permissionLevel
        for membership in clinic_memberships
        if normalize_permission_level(membership.permissionLevel) != "read"
    ]
    if invalid_permissions:
        raise HTTPException(
            status_code=403,
            detail="Workstation accounts must be read-only",
        )

    organisation_id = clinic_memberships[0].organisationId

    user.lastLoginAt = datetime.utcnow()
    db.commit()

    access_token = create_access_token(
        user_id=user.userId,
        organisation_id=organisation_id,
        account_type=account_type,
    )

    return LoginResponse(
        accessToken=access_token,
        mustSetPassword=bool(user.mustSetPassword),
        user=AuthUserPayload(
            userId=user.userId,
            displayName=user.displayName,
            email=user.email,
            username=getattr(user, "username", None),
            accountType=account_type,
            organisationId=organisation_id,
            organisationPermissionLevel=None,
            clinicMemberships=[
                ClinicMembershipPayload(
                    clinicId=membership.clinicId,
                    permissionLevel=normalize_permission_level(membership.permissionLevel),
                )
                for membership in clinic_memberships
            ],
        ),
    )

@app.post("/auth/change-password")
def change_password(
    req: ChangePasswordRequest,
    db: Session = Depends(get_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    user = get_user_by_id(db, current_user.userId)
    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    if not user.passwordHash:
        raise HTTPException(
            status_code=400,
            detail="No current password is set for this account",
        )

    if not verify_password(req.currentPassword, user.passwordHash):
        raise HTTPException(status_code=400, detail="Current password is incorrect")

    validate_password_strength(req.newPassword)

    user.passwordHash = hash_password(req.newPassword)
    user.mustSetPassword = False
    db.commit()

    return {"status": "password_changed"}

@app.post("/auth/reset-password")
def reset_password(
    req: ResetPasswordRequest,
    db: Session = Depends(get_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    target_user = get_user_by_id(db, req.targetUserId)
    if not target_user:
        raise HTTPException(status_code=404, detail="User not found")

    require_password_reset_access(
        db,
        actor=current_user,
        target_user=target_user,
    )

    target_user.passwordHash = None
    target_user.mustSetPassword = True
    db.commit()

    return {"status": "password_reset_required"}

@app.get("/network/settings")
def get_network_settings(
    request: Request,
    current_user: CurrentUser = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    clinic_id = current_user.selectedClinicId
    organisation_id = current_user.organisationId

    if not clinic_id:
        raise HTTPException(status_code=400, detail="No clinic selected")

    # Get network mode
    network = (
        db.query(NetworkAccess)
        .filter(
            NetworkAccess.organisationId == organisation_id,
            NetworkAccess.clinicId == clinic_id,
        )
        .first()
    )

    mode = "public"
    if network:
        mode = (network.mode or "public").lower()

    # Get allowed IPs
    allowed_ip_rows = (
        db.query(AllowedIP)
        .filter(
            AllowedIP.organisationId == organisation_id,
            AllowedIP.clinicId == clinic_id,
        )
        .all()
    )

    entries = [
        {
            "id": row.allowedIpId,
            "value": row.value,
            "label": row.label,
        }
        for row in allowed_ip_rows
    ]

    # Detect current IP
    current_ip = get_client_ip(request)

    # Check if already allowed
    current_ip_already_added = any(
        (row.value or "").strip() == current_ip for row in allowed_ip_rows
    )

    return {
        "clinicId": clinic_id,
        "mode": mode,
        "currentIp": current_ip,
        "entries": entries,
        "currentIpAlreadyAdded": current_ip_already_added,
    }

@app.post("/network/settings")
def save_network_settings(
    req: SaveNetworkSettingsRequest,
    current_user: CurrentUser = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    clinic_id = current_user.selectedClinicId
    organisation_id = current_user.organisationId

    if not clinic_id:
        raise HTTPException(status_code=400, detail="No clinic selected")

    effective_level = get_effective_permission_level(
        db,
        user=current_user,
        clinic_id=clinic_id,
    )

    if not effective_level or not has_level(effective_level, "manage"):
        raise HTTPException(
            status_code=403,
            detail="You do not have permission to update network settings",
        )

    normalized_mode = (req.mode or "").strip().lower()
    if normalized_mode not in {"public", "restricted"}:
        raise HTTPException(status_code=400, detail="Invalid network mode")

    normalized_entries = []
    seen_values = set()

    for item in req.entries or []:
        value = str(item.value or "").strip()
        label = str(item.label or "").strip() or None

        if not value:
            raise HTTPException(status_code=400, detail="IP value is required")

        if value in seen_values:
            raise HTTPException(
                status_code=400,
                detail=f"Duplicate IP entry: {value}",
            )

        seen_values.add(value)
        normalized_entries.append(
            {
                "value": value,
                "label": label,
            }
        )

    if normalized_mode == "restricted" and not normalized_entries:
        raise HTTPException(
            status_code=400,
            detail="At least one allowed IP is required when restricted mode is enabled",
        )

    network = (
        db.query(NetworkAccess)
        .filter(
            NetworkAccess.organisationId == organisation_id,
            NetworkAccess.clinicId == clinic_id,
        )
        .first()
    )

    if not network:
        network = NetworkAccess(
            networkAccessId=str(uuid.uuid4()),
            organisationId=organisation_id,
            clinicId=clinic_id,
            mode=normalized_mode,
        )
        db.add(network)
    else:
        network.mode = normalized_mode
        network.updatedAt = datetime.utcnow()

    existing_allowed_ips = (
        db.query(AllowedIP)
        .filter(
            AllowedIP.organisationId == organisation_id,
            AllowedIP.clinicId == clinic_id,
        )
        .all()
    )

    for row in existing_allowed_ips:
        db.delete(row)

    for item in normalized_entries:
        db.add(
            AllowedIP(
                allowedIpId=str(uuid.uuid4()),
                organisationId=organisation_id,
                clinicId=clinic_id,
                value=item["value"],
                label=item["label"],
            )
        )

    db.commit()

    return {
        "status": "saved",
        "clinicId": clinic_id,
        "mode": normalized_mode,
        "entries": normalized_entries,
    }